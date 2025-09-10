import os
import sys
import shutil
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
import pickle
import json
import locale
from datetime import datetime, timedelta, time
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
#import re
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_BREAK

# Variable global para reutilizar la ruta
ruta_excel_principal = None
ruta_carpeta_global = None
ruta_plantillas = None
seleccion_primera_vez = True

# DEBUG flag (pon True solo mientras se depura)
DEBUG_LEER_REPORTE = True

# -------------------------------------------------------------------------------
# Utilidades generales (sin cambios esenciales)
# -------------------------------------------------------------------------------

def obtener_iniciales(nombre_completo: str) -> str:
    partes = str(nombre_completo).split()
    return ''.join([p[0].upper() for p in partes if p])

try:
    locale.setlocale(locale.LC_TIME, "es_ES.UTF-8")
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, "Spanish_Spain.1252")
    except locale.Error:
        pass

CONFIG_FILE = "config.json"


def guardar_config(carpeta_raiz):
    with open(CONFIG_FILE, "w") as f:
        json.dump({"carpeta_raiz": carpeta_raiz}, f)


def cargar_config():
    try:
        with open(CONFIG_FILE, "r") as f:
            return json.load(f).get("carpeta_raiz", "")
    except FileNotFoundError:
        return ""


# -------------------------------------------------------------------------------
# NUEVO: Utilidades específicas para agregar actividad y evidencia
# -------------------------------------------------------------------------------

def previsualizar_tabla_debug(df, titulo="Previsualización - primeras filas"):
    """Muestra una ventana con las primeras filas (debug), sin bloquear toda la app."""
    try:
        top = tk.Toplevel(root)
        top.title(titulo)
        top.geometry("700x300+100+100")  # tamaño fijo, con coordenadas para evitar que quede fuera de pantalla

        # Texto scrollable
        txt = tk.Text(top, wrap="none", font=("Consolas", 10))
        txt.pack(fill="both", expand=True, padx=5, pady=5)

        scrollbar_v = ttk.Scrollbar(top, orient="vertical", command=txt.yview)
        scrollbar_v.pack(side="right", fill="y")
        txt.configure(yscrollcommand=scrollbar_v.set)

        # Mostrar primeras filas
        txt.insert("1.0", df.head(20).to_string(index=False))

        # Botón para cerrar
        tk.Button(top, text="Continuar", command=top.destroy, bg="#4CAF50", fg="white", font=("Arial", 12, "bold")).pack(pady=6)

        # 🔹 YA NO usamos grab_set ni wait_window → así no se queda colgado
        # top.transient(root)  # opcional: para que quede arriba de root
    except Exception as e:
        print("No se pudo abrir la vista debug:", e)


def _map_nombre_columna(raw):
    """Mapea nombres de columna detectados a los nombres esperados de manera tolerante."""
    if raw is None:
        return raw
    r = str(raw).strip().lower()
    # coincidencias por palabra clave
    if "nombre" in r and "apellidos" not in r and "apellido" not in r:
        return "Nombre"
    if "apellido" in r:
        return "Apellidos"
    if "fecha" in r and "atenc" in r:
        return "Fecha de atención"
    if "fecha" in r and "atenc" not in r:
        # aceptamos 'Fecha' como 'Fecha de atención'
        return "Fecha de atención"
    if "tema" in r or "asunto" in r:
        return "Tema o asunto tratado"
    if "situac" in r or "situa" in r:
        return "Situación del alumno"
    # por defecto: devolver título "limpio" (capitalizado)
    return str(raw).strip()

def normalizar_columnas(df: pd.DataFrame) -> pd.DataFrame:
    """Limpia espacios en los encabezados y estandariza nombres esperados."""
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    # Aliases tolerantes
    renombres = {
        "Tema o asunto tratado": "Tema o asunto tratado",
        "Tema o asunto tratado ": "Tema o asunto tratado",
        "Situación del alumno": "Situación del alumno",
        "Fecha de atención": "Fecha de atención",
        "Nombre": "Nombre",
        "Apellidos": "Apellidos",
    }
    df.rename(columns=renombres, inplace=True)
    return df


def seleccionar_excel_tabla_reporte() -> str:
    global ruta_excel_principal
    #messagebox.showinfo("Instrucción", "Selecciona el Excel semanal que contiene la tabla 'TablaReporte' (hoja 3)")
    #rutaExc = rutaExc
    #ruta = rutaExc
    #ruta = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx;*.xls;*.xlsm")], title="Selecciona Excel semanal")
    if not ruta_excel_principal:
        raise SystemExit("No se seleccionó Excel semanal.")
    return ruta_excel_principal


def leer_tabla_reporte(ruta_excel_principal: str, DEBUG_LEER_REPORTE=False) -> pd.DataFrame:
    """
    Lee la hoja 3 (índice 2) del Excel y devuelve un DataFrame filtrado de la tabla de reporte.
    Detecta automáticamente el encabezado en B11:J11 y valida que existan las columnas 'Nombre' y 'Apellidos'.
    """
    try:
        # Leer la hoja completa
        df_full = pd.read_excel(ruta_excel_principal, sheet_name=2, header=None)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo leer la hoja 3 del Excel. Detalle: {e}")
        raise

    # Tomar solo las columnas B:J y filas desde la fila 11 (índice 10) para detectar encabezado
    df_header = df_full.iloc[10, 1:10]  # B11:J11
    df_data = df_full.iloc[11:, 1:10]   # Datos debajo del encabezado
    df_data.columns = [str(c).strip() for c in df_header]

    # --- Validación de columnas críticas ---
    cols_lower = [c.lower() for c in df_data.columns]
    tiene_nombre = any(c == "nombre" for c in cols_lower)
    tiene_apellidos = any(c == "apellidos" or c == "apellido" for c in cols_lower)
    if not (tiene_nombre and tiene_apellidos):
        cols_found = ", ".join(df_data.columns.tolist())
        raise ValueError(
            f"La hoja 3 no contiene las columnas esperadas 'Nombre' y 'Apellidos'.\n"
            f"Columnas detectadas: {cols_found}"
        )

    # Generar columna de nombre completo (Reporte)
    df_data["Nombre completo (Reporte)"] = (
        df_data["Nombre"].astype(str).str.strip() + " " +
        df_data["Apellidos"].astype(str).str.strip()
    ).str.strip()

    # Opción de debug: mostrar primeras filas
    if DEBUG_LEER_REPORTE:
        try:
            previsualizar_tabla_debug(df_data, titulo="DEBUG_LEER_REPORTE")
        except Exception as e:
            print("No se pudo abrir la ventana de debug:", e)

    return df_data.reset_index(drop=True)




def filtrar_fila_reciente_por_alumno(df_reporte: pd.DataFrame, nombre_objetivo: str) -> pd.Series:
    """
    Selecciona la fila más reciente del alumno en df_reporte comparando
    'Nombre completo (Reporte)'. Si no encuentra, lanza error con pistas.
    """
    if "Nombre completo (Reporte)" not in df_reporte.columns:
        # intentar reconstruir desde columnas detectables
        cols = df_reporte.columns.tolist()
        if any("nombre" in c.lower() for c in cols) and any("apell" in c.lower() for c in cols):
            ncol = next(c for c in cols if "nombre" in c.lower())
            acol = next(c for c in cols if "apell" in c.lower())
            df_reporte["Nombre completo (Reporte)"] = (df_reporte[ncol].fillna("").astype(str).str.strip() + " " + df_reporte[acol].fillna("").astype(str).str.strip()).str.strip()
        else:
            raise ValueError("No se pudo construir 'Nombre completo (Reporte)' desde las columnas existentes.")

    # Comparación insensible a mayúsculas y espacios
    objetivo_norm = str(nombre_objetivo).strip().upper()
    df_reporte["___NOM_SEARCH__"] = df_reporte["Nombre completo (Reporte)"].astype(str).str.strip().str.upper()

    df_al = df_reporte[df_reporte["___NOM_SEARCH__"] == objetivo_norm].copy()
    if df_al.empty:
        # Preparar mensaje útil para depurar: listar algunos nombres detectados
        uniques = df_reporte["Nombre completo (Reporte)"].dropna().unique().tolist()[:30]
        raise ValueError(f"No se encontraron registros para '{nombre_objetivo}' en TablaReporte. Algunos nombres detectados: {uniques[:10]} ...")
    # ordenar por Fecha de atención si existe
    if "Fecha de atención" in df_al.columns:
        df_al["Fecha de atención"] = pd.to_datetime(df_al["Fecha de atención"], errors="coerce")
        df_al.sort_values(by="Fecha de atención", inplace=True, na_position="first")
        fila = df_al.iloc[-1]
    else:
        fila = df_al.iloc[-1]
    # limpiar columna temporal
    df_reporte.drop(columns=["___NOM_SEARCH__"], errors="ignore", inplace=True)
    return fila


def buscar_carpeta_alumno(ruta_carpeta_global: str, matricula: str, nombre_mayus: str, programa: str) -> str:
    base_folder = f"{matricula}_{nombre_mayus}_{programa}"
    ruta = os.path.join(ruta_carpeta_global, base_folder)
    if not os.path.isdir(ruta):
        raise FileNotFoundError(f"No se encontró la carpeta del alumno: {base_folder}")
    return ruta


def encontrar_archivo_seguimiento(ruta_alumno: str) -> str:
    """Busca el archivo Word que comienza con '5_Seguimiento individual_' en el árbol del alumno."""
    for root_dir, _, files in os.walk(ruta_alumno):
        for f in files:
            if f.lower().endswith(".docx") and f.startswith("5_Seguimiento individual_"):
                return os.path.join(root_dir, f)
    raise FileNotFoundError("No se encontró un archivo que comience con '5_Seguimiento individual_'.")

def encontrar_archivo_entrevista(ruta_alumno: str) -> str:
    """Busca el archivo Word que comienza con '3_Entrevista extendida_' en el árbol del alumno."""
    for root_dir, _, files in os.walk(ruta_alumno):
        for f in files:
            if f.lower().endswith(".docx") and f.startswith("3_Entrevista extendida_"):
                return os.path.join(root_dir, f)
    raise FileNotFoundError("No se encontró un archivo que comience con '3_Entrevista extendida_'.")


def encontrar_carpeta_evidencias(ruta_alumno: str) -> str:
    for root_dir, dirs, _ in os.walk(ruta_alumno):
        for d in dirs:
            if d.startswith("6_Evidencias_"):
                return os.path.join(root_dir, d)
    raise FileNotFoundError("No se encontró carpeta que comience con '6_Evidencias_'.")


def primera_fila_col4_vacia(tabla: docx.table.Table) -> int:
    """Devuelve el índice de fila (0-based) donde la columna 4 esté vacía por primera vez.
    Si no hay, devuelve len(rows) para que podamos anexar una nueva fila."""
    for i, row in enumerate(tabla.rows):
        celdas = row.cells
        if len(celdas) >= 4:
            texto_col4 = celdas[3].text.strip()
            if texto_col4 == "":
                return i
    return len(tabla.rows)


def asegurar_filas(tabla: docx.table.Table, filas_necesarias: int):
    """Agrega filas a la tabla hasta alcanzar el total filas_necesarias (si hace falta)."""
    while len(tabla.rows) < filas_necesarias:
        tabla.add_row()


def escribir_en_fila(tabla: docx.table.Table, idx: int, col2: str, col3: str, col4: str):
    asegurar_filas(tabla, idx + 1)
    row = tabla.rows[idx]
    # Asegurar 4 columnas
    if len(row.cells) < 4:
        raise ValueError("La tabla de seguimiento debe tener al menos 4 columnas.")
    # Col 2 (Fecha), Col 3 (Tema), Col 4 (Situación)
    row.cells[1].text = str(col2) if col2 is not None else ""
    row.cells[2].text = str(col3) if col3 is not None else ""
    row.cells[3].text = str(col4) if col4 is not None else ""
    # Opcional: tamaño de fuente uniforme
    for c in [row.cells[1], row.cells[2], row.cells[3]]:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)


import docx
from docx.shared import Pt

import tkinter as tk
from tkinter import messagebox

def seleccionar_tutor_y_asistencia(nombre):
    """
    Muestra ventana para seleccionar tutor y si asistió el alumno.
    Retorna:
        nombre_tutor (str)
        asistio (bool)
    """
    ventana = tk.Toplevel(root)
    ventana.title("Información del Tutor y Asistencia para "+ str(nombre)) #AAAAA
    #ventana.geometry("1000x300")
    ventana.resizable(False, False)

    # Tamaño fijo y centrado
    ancho, alto = 1000, 300
    x = (ventana.winfo_screenwidth() // 2) - (ancho // 2)
    y = (ventana.winfo_screenheight() // 2) - (alto // 2)
    ventana.geometry(f"{ancho}x{alto}+{x}+{y}")
    ventana.resizable(False, False)
    ventana.protocol("WM_DELETE_WINDOW", ventana.destroy)
    
    nombre_tutor = tk.StringVar(value="")
    asistio = tk.StringVar(value="")

    tk.Label(ventana, text="Selecciona el tutor para "+ nombre + ":", font=("Arial", 12, "bold")).pack(pady=10)

    # Frame para botones de tutores
    frame_tutores = tk.Frame(ventana)
    frame_tutores.pack(pady=5)

    tutores = ["Diana Lucia Herrera Guízar", "Luis Alberto Ramos Llanos", "Yenifer Fuerte Cortés", "Alan Alvarado Ramírez", "Veda López Báez"]

    botones_tutores = {}

    def marcar_tutor(tutor):
        nombre_tutor.set(tutor)
        # Resetear colores
        for t, btn in botones_tutores.items():
            btn.config(bg="SystemButtonFace")
        # Marcar seleccionado
        botones_tutores[tutor].config(bg="lightgreen")

    for t in tutores:
        btn = tk.Button(frame_tutores, text=t, width=20, command=lambda t=t: marcar_tutor(t))
        btn.pack(side="left", padx=5, pady=5)
        botones_tutores[t] = btn

    # Pregunta asistencia
    tk.Label(ventana, text="¿Asistió el alumno?", font=("Arial", 12)).pack(pady=(20,5)) #AAAAA
    frame_asistencia = tk.Frame(ventana)
    frame_asistencia.pack(pady=5)

    def marcar_asistencia(valor):
        asistio.set(valor)
        ventana.destroy()  # cerrar ventana una vez tomada la decisión

    tk.Button(frame_asistencia, text="Sí", width=10, bg="lightblue",
              command=lambda: marcar_asistencia("Sí")).pack(side="left", padx=20)
    tk.Button(frame_asistencia, text="No", width=10, bg="lightcoral",
              command=lambda: marcar_asistencia("No")).pack(side="left", padx=20)

    # Esperar a que se cierre la ventana
    ventana.grab_set()
    ventana.wait_window()

    if not nombre_tutor.get() or asistio.get() is None:
        messagebox.showwarning("Atención", "Debes seleccionar un tutor y la asistencia del alumno.")
        return seleccionar_tutor_y_asistencia()  # Vuelve a preguntar

    return nombre_tutor.get(), asistio.get()


def insertar_actividad_en_word(ruta_docx: str, fila: pd.Series, nombre):
    """
    Inserta los datos de una fila de reporte en la tabla del Word.
    - Detecta automáticamente la primera fila vacía en la columna 4.
    - Convierte fechas a string.
    - Acepta un pandas.Series (fila del reporte).
    """
    if not os.path.isfile(ruta_docx):
        raise FileNotFoundError(f"No se encontró el archivo Word: {ruta_docx}")

    try:
        doc = docx.Document(ruta_docx)
    except Exception as e:
        raise RuntimeError(f"No se pudo abrir el documento Word: {e}")

    # Solicitar tutor y asistencia antes de insertar
    nombre_tutor, asistio = seleccionar_tutor_y_asistencia(nombre)

    # Buscar la primera tabla (o la que necesites)
    if not doc.tables:
        raise ValueError(f"El documento {ruta_docx} no contiene tablas.")

    tabla = doc.tables[0]  # <-- Ajustar si hay más de una tabla
    # Columnas que necesitamos insertar: Fecha (2), Tema (3), Situación (4)
    col_fecha = 1
    col_tema = 2
    col_situacion = 3
    col_Asistencia = 4
    col_Tutor = 5


    # Convertir valores a string seguro
    fecha = fila.get("Fecha de atención")
    Hora = str(fila.get("Hora"))

    if pd.notna(fecha):
        if isinstance(fecha, pd.Timestamp):
            fecha = fecha.strftime("%d/%m/%Y") + " " + Hora
        else:
            fecha = str(fecha).strip()
    else:
        fecha = ""

    tema = str(fila.get("Tema o asunto tratado", "")).strip()
    situacion = str(fila.get("Situación del alumno", "")).strip()

    # Buscar primera fila vacía en la columna 4 (Situación)
    fila_insert = None
    for row in tabla.rows:
        celda = row.cells[col_situacion]
        if not celda.text.strip():
            fila_insert = row
            break
    if fila_insert is None:
        # Si no hay fila vacía, agregar una nueva fila
        fila_insert = tabla.add_row()

    # Insertar valores en la fila encontrada
    fila_insert.cells[col_fecha].text = fecha
    fila_insert.cells[col_tema].text = tema
    fila_insert.cells[col_situacion].text = situacion
    fila_insert.cells[col_Asistencia].text = str(asistio)
    fila_insert.cells[col_Tutor].text = str(nombre_tutor)

    # Alinear texto al centro y ajustar fuente
    for cell in [fila_insert.cells[col_fecha], fila_insert.cells[col_tema],
                fila_insert.cells[col_situacion], fila_insert.cells[col_Asistencia],
                fila_insert.cells[col_Tutor]]:
        
        # Alineación de celda (vertical)
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER  # Esto centra verticalmente
        
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centra horizontalmente
            for run in p.runs:
                run.font.size = Pt(10)  # Ajustar tamaño según necesites
    
    # Guardar cambios
    doc.save(ruta_docx)



def proximo_numero_sesion(carpeta_evidencias: str, iniciales: str) -> int:
    """Calcula el siguiente número de sesión buscando archivos existentes 'sesión_#_..._INICIALES'."""
    extensiones = (".jpg", ".jpeg", ".png")
    archivos = [f for f in os.listdir(carpeta_evidencias) if f.lower().endswith(extensiones)]
    return len(archivos) + 1


def copiar_y_renombrar_evidencia(carpeta_evidencias: str, ruta_imagen_src: str, actividad: str, iniciales: str) -> str:
    num_sesion = proximo_numero_sesion(carpeta_evidencias, iniciales)
    # Limpiar nombre de actividad para archivo
    actividad_limpia = "_".join(str(actividad).strip().split())
    nombre_dest = f"Sesión{num_sesion}_{actividad_limpia}_{iniciales}{os.path.splitext(ruta_imagen_src)[1].lower()}"
    ruta_dest = os.path.join(carpeta_evidencias, nombre_dest)
    shutil.copy2(ruta_imagen_src, ruta_dest)
    return ruta_dest


# -------------------------------------------------------------------------------
# Módulo crear_expedientes (MODIFICADO mínimamente)
# -------------------------------------------------------------------------------

def crear_expedientes(datos):

    global seleccion_primera_vez
    global ruta_carpeta_global
    global ruta_plantillas

    if seleccion_primera_vez:
        ruta_carpeta_global = seleccionar_carpeta_raiz("Selecciona la carpeta donde se crearán los expedientes:")
        #ruta_carpeta_global = destino
        plantilla_dir = seleccionar_carpeta("Selecciona la carpeta donde están las plantillas:")
        ruta_plantillas = plantilla_dir
        seleccion_primera_vez = False

    prefijo = simpledialog.askstring("Prefijo de Carpeta", "Introduce el prefijo antes de '@' en la carpeta interna:")

    if not prefijo:
        messagebox.showerror("Error", "No ingresaste un prefijo. Proceso cancelado.")
        return

    for i in range(len(datos.get("Nombre completo", []))):
        nombre = datos["Nombre completo"][i].upper()
        matricula = datos.get("Matrícula", [None])[i]
        programa = datos.get("Programa", [None])[i]
        periodo = datos.get("Período", [None])[i]

        if not (nombre and matricula and programa and periodo):
            continue

        base_folder = f"{matricula}_{nombre}_{programa}"
        ruta_alumno = os.path.join(ruta_carpeta_global, base_folder)
        os.makedirs(ruta_alumno, exist_ok=True)

        iniciales = obtener_iniciales(nombre)
        interna = f"{prefijo}@{periodo}_{matricula}_{iniciales}_{programa}"
        ruta_interna = os.path.join(ruta_alumno, interna)
        os.makedirs(ruta_interna, exist_ok=True)

        # Copiar/renombrar plantillas y carpetas
        for item in os.listdir(plantilla_dir):
            src_path = os.path.join(plantilla_dir, item)
            if os.path.isfile(src_path):
                nuevo_nombre = f"{os.path.splitext(item)[0]}_{iniciales}_{programa}_{periodo}{os.path.splitext(item)[1]}"
                dest_path = os.path.join(ruta_interna, nuevo_nombre)
                shutil.copy2(src_path, dest_path)
            elif os.path.isdir(src_path):
                nuevo_nombre_carpeta = f"{item}_{iniciales}_{programa}_{periodo}"
                dest_dir_path = os.path.join(ruta_interna, nuevo_nombre_carpeta)
                os.makedirs(dest_dir_path, exist_ok=True)
                for archivo_interno in os.listdir(src_path):
                    src_file_path = os.path.join(src_path, archivo_interno)
                    if archivo_interno.lower().endswith((".jpeg", ".jpg")):
                        nombre_base, extension = os.path.splitext(archivo_interno)
                        nuevo_nombre_archivo = f"{nombre_base}_{iniciales}{extension}"
                    else:
                        nuevo_nombre_archivo = archivo_interno
                    dest_file_path = os.path.join(dest_dir_path, nuevo_nombre_archivo)
                    shutil.copy2(src_file_path, dest_file_path)

    messagebox.showinfo("Proceso Completo", "Los expedientes fueron creados exitosamente.")
    mostrar_crear_opciones(datos)


def mostrar_crear_opciones(datos):
    win = tk.Toplevel(root)
    win.title("Opciones después de crear expedientes")

    # Tamaño fijo y centrado en pantalla
    ancho, alto = 320, 200
    x = (win.winfo_screenwidth() // 2) - (ancho // 2)
    y = (win.winfo_screenheight() // 2) - (alto // 2)
    win.geometry(f"{ancho}x{alto}+{x}+{y}")
    win.resizable(False, False)

    #win.geometry("320x200")
    tk.Label(win, text="¿Qué deseas hacer ahora?").pack(pady=10)
    tk.Button(
        win,
        text="Llenar Documentos",
        command=lambda: (llenar_documentos(datos, win), win.destroy())
    ).pack(pady=5)
    tk.Button(win, text="Continuar sin llenar", command=win.destroy).pack(pady=5)


# -------------------------------------------------------------------------------
# Módulo llenar_documentos (SIN CAMBIOS funcionales)
# -------------------------------------------------------------------------------

def llenar_documentos(datos, ventana=None):
    raiz = cargar_config()
    if not raiz:
        messagebox.showerror("Error", "No se encontró carpeta raíz. Ejecuta primero crear_expedientes.")
        return

    NombreTutor = simpledialog.askstring("Nombre del tutor", "Introduce el nombre del Tutor:")

    for i in range(len(datos.get("Nombre completo", []))):
        nombre = datos["Nombre completo"][i]
        matricula = datos["Matrícula"][i]
        programa = datos["Programa"][i]
        periodo = datos["Período"][i]
        dia = datos.get("Dia", [None]*len(datos["Nombre completo"]))[i]
        hora_ini_raw = datos.get("Hora inicio", [None]*len(datos["Nombre completo"]))[i]
        horafin_raw = datos.get("Hora final", [None]*len(datos["Nombre completo"]))[i]

        hora_ini = formatear_hora(hora_ini_raw)
        horafin = formatear_hora(horafin_raw)

        horario = f"{dia} {hora_ini}-{horafin}"
        area = datos.get("Área de intervención", [None])[i] or ""
        edad = datos.get("Edad", [None])[i] or ""
        semestre = datos.get("Semestre", [None])[i] or ""
        habilidades = datos.get("Habilidades", [None])[i] or ""
        grupoProv = datos.get("Grupo", [None])[i] or ""
        titularDir = datos.get("Titular/Director", [None])[i] or ""
        situacion = datos.get("Situación", [None])[i] or ""
        promedio = datos.get("Promedio Anterior", [None])[i] or ""

        base_folder = f"{matricula}_{nombre.upper()}_{programa}"
        ruta_base = os.path.join(raiz, base_folder)

        programa_bach = ""

        if programa == "BUNLA" or programa == "BIUNLA" or programa == "BUNLAV":
            programa_bach = programa
            programa = ""
        else:
            grupoProv = ""

        if not os.path.isdir(ruta_base):
            print(f"Advertencia: No se encontró la carpeta para {nombre}. Saltando...")
            continue

        for root_dir, _, files in os.walk(ruta_base):
            for file in files:
                if not file.lower().endswith(".docx"):
                    continue

                ruta_doc = os.path.join(root_dir, file)
                try:
                    doc = docx.Document(ruta_doc)

                    markers = {
                        "{NOMBRE}": nombre,
                        "{NOMBREMAY}": nombre.upper(),
                        "{MATRICULA}": matricula,
                        "{PROGRAMA}": programa,
                        "{PROGRAMABACH}": programa_bach,
                        "{GRUPO}": grupoProv,
                        "{TitularDir}": titularDir,
                        "{SITUACIÓN}": situacion,
                        "{PERIODO}": periodo,
                        "{PROMEDIO}": promedio,
                        "{EDAD}": str(edad)+" años",
                        "{AREA}": area,
                        "{TUTOR}": NombreTutor,
                        "{SEMESTRE}": str(semestre),
                        "{HORARIO}": horario,
                        "{HABILIDADES}": habilidades,
                        "{FECHA}": datetime.now().strftime("%d/%m/%Y"),
                        "{FECHALARGA}": datetime.now().strftime("%d de %B de %Y"),
                        "{FECHADIASNANO}": datetime.now().strftime("%A %d de %B").upper(),
                        "{HORAINICIO}": hora_ini,
                        "{HORAFINAL}": horafin,
                        "{DIAINICIO}": dia,
                        "{PROXREUNION}": proxima_reunion(dia, hora_ini)
                    }

                    reemplazar_en_docx(doc, markers)
                    doc.save(ruta_doc)
                except Exception as e:
                    print(f"No se pudo procesar el archivo {ruta_doc}. Error: {e}")

    messagebox.showinfo("Proceso Completo", "Los documentos fueron llenados exitosamente.")


# -------------------------------------------------------------------------------
# Resto de utilidades existentes
# -------------------------------------------------------------------------------

def enviar_bienvenida(datos):
    print("[Correo] Bienvenida:", datos)


def enviar_extranamiento(datos):
    print("[Correo] Extranamiento:", datos)


def añadir_a_pit(datos):
    pass  # TODO: implementar lógica PIT


def imprimir_entrevista_b1(): print("Imprimiendo Entrevista B1...")

def imprimir_entrevista_bcont(): print("Imprimiendo Entrevista BCont...")

def imprimir_entrevista_l1(): print("Imprimiendo Entrevista L1...")

def imprimir_entrevista_lcont(): print("Imprimiendo Entrevista LCont...")

def imprimir_reglas(): print("Imprimiendo Reglas...")


def seleccionar_carpeta(mensaje):
    messagebox.showinfo("Instrucción", mensaje)
    ruta = filedialog.askdirectory(title=mensaje)
    if not ruta:
        sys.exit(0)
    return ruta


def seleccionar_carpeta_raiz(mensaje):
    messagebox.showinfo("Instrucción", mensaje)
    ruta = filedialog.askdirectory(title=mensaje)
    if not ruta:
        sys.exit(0)
    guardar_config(ruta)
    return ruta


def cargar_datos_excel():
    global ruta_excel_principal
    messagebox.showinfo("Instrucción", "Selecciona el archivo Excel con datos de los alumnos")
    ruta = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx;*.xls;*.xlsm")])
    ruta_excel_principal = ruta
    if not ruta:
        messagebox.showwarning("Aviso", "No se seleccionó archivo.")
        return None

    # Limpieza y tipos
    df = pd.read_excel(ruta, dtype={"Matrícula": str, "Grupo": str})
    df.dropna(subset=['Nombre completo'], inplace=True)
    df = df[df['Nombre completo'].astype(str).str.strip().str.lower() != 'nan']
    df = df[df['Nombre completo'].astype(str).str.strip() != '']

    if df.empty:
        messagebox.showwarning("Aviso", "No se encontraron alumnos con nombre válido en el archivo.")
        return None

    datos = {col: df[col].astype(str).tolist() for col in df.columns}

    with open('datos.pkl', 'wb') as f:
        pickle.dump(datos, f)

    mostrar_alumnos(datos)
    return datos


def filtrar_datos(datos, matricula):
    datos_filtrados = {col: [] for col in datos.keys()}
    for i, m in enumerate(datos['Matrícula']):
        if m.strip() == matricula.strip():
            for col in datos:
                datos_filtrados[col].append(datos[col][i])
    return datos_filtrados


def cerrar_programa():
    sys.exit()


def mostrar_alumnos(datos):
    win = tk.Toplevel(root)
    win.title("Seleccionar Alumno(s)")

    # Tamaño fijo y centrado
    ancho, alto = 410, 600
    x = (win.winfo_screenwidth() // 2) - (ancho // 2)
    y = (win.winfo_screenheight() // 2) - (alto // 2)
    win.geometry(f"{ancho}x{alto}+{x}+{y}")
    win.resizable(False, False)
    win.protocol("WM_DELETE_WINDOW", cerrar_programa)

    # Encabezado
    tk.Label(
        win,
        text="Seleccione uno o varios alumnos:",
        font=("Arial", 13, "bold"),
        pady=10
    ).pack()

    # Contenedor con scroll
    frame = tk.Frame(win)
    frame.pack(fill="both", expand=True, padx=10, pady=5)

    canvas = tk.Canvas(frame, borderwidth=0)
    scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas)

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", width=350)
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Habilitar scroll con mouse
    def _on_mousewheel(event):
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    canvas.bind_all("<MouseWheel>", _on_mousewheel)

    # Lista de alumnos con checkboxes
    vars_checks = []
    for i, nombre in enumerate(datos['Nombre completo']):
        var = tk.BooleanVar(value=False)
        chk = tk.Checkbutton(
            scrollable_frame,
            text=nombre,
            variable=var,
            font=("Arial", 11)
        )
        chk.pack(anchor="w", padx=10, pady=2)
        vars_checks.append((var, i))

    # Procesar selección
    def procesar_seleccion():
        seleccionados = [i for var, i in vars_checks if var.get()]
        if not seleccionados:
            messagebox.showwarning("Atención", "Selecciona al menos un alumno.")
            return
        datos_seleccionados = {col: [] for col in datos.keys()}
        for col in datos.keys():
            for idx in seleccionados:
                datos_seleccionados[col].append(datos[col][idx])
        ventana_opciones(datos_seleccionados)

    # Botones de acción
    botones_frame = tk.Frame(win)
    botones_frame.pack(pady=15)

    tk.Button(
        botones_frame,
        text="✅ Procesar selección",
        command=procesar_seleccion,
        font=("Arial", 11, "bold"),
        bg="#4CAF50", fg="white",
        activebackground="#388E3C",
        padx=15, pady=8
    ).pack(side="left", padx=10)

    tk.Button(
        botones_frame,
        text="☑ Seleccionar Todos",
        command=lambda: seleccionar_todos(datos, win),
        font=("Arial", 11, "bold"),
        bg="#2196F3", fg="white",
        activebackground="#1976D2",
        padx=15, pady=8
    ).pack(side="left", padx=10)




def seleccionar_todos(datos, win=None):
    if win:
        pass
    ventana_opciones(datos)


def ventana_opciones(datos):
    """
    Ventana de opciones que procesa todos los alumnos seleccionados
    de forma secuencial al presionar los botones.
    """
    win = tk.Toplevel(root)
    win.title("Opciones")

    ancho, alto = 400, 330
    x = (win.winfo_screenwidth() // 2) - (ancho // 2)
    y = (win.winfo_screenheight() // 2) - (alto // 2)
    win.geometry(f"{ancho}x{alto}+{x}+{y}")
    win.resizable(False, False)

    #if len(datos['Nombre completo']) > 1:
    #    opcion = 2
    #else:
    #    opcion = 1

    tk.Label(
        win,
        text=f"Opciones para {'todos los alumnos' if len(datos['Nombre completo']) > 1 else datos['Nombre completo'][0]}",
        font=("Arial", 10, "bold"),
        pady=15
    ).pack()

    # Crear expedientes para todos
    tk.Button(
        win,
        text="📂 Crear Expedientes",
        command=lambda: (crear_expedientes(datos), win.destroy()),
        font=("Arial", 12, "bold"),
        bg="#2196F3", fg="white",
        activebackground="#1976D2",
        padx=20, pady=10
    ).pack(pady=10)

    # Función auxiliar para procesar cada alumno uno por uno
    def procesar_actividad():
        for idx in range(len(datos["Nombre completo"])):
            alumno = {col: [datos[col][idx]] for col in datos.keys()}
            agregar_actividad_y_evidencia(alumno)
        win.destroy()

    def procesar_entrevista():
        for idx in range(len(datos["Nombre completo"])):
            alumno = {col: [datos[col][idx]] for col in datos.keys()}
            agregar_entrevista_extendida(alumno)
        win.destroy()

    # Botones de actividad y entrevista
    tk.Button(
        win,
        text="📝 Agregar actividad y evidencia",
        command=lambda: procesar_actividad(),
        font=("Arial", 11, "bold"),
        bg="#4CAF50", fg="white",
        activebackground="#388E3C",
        padx=15, pady=8
    ).pack(pady=8)

    tk.Button(
        win,
        text="💬 Agregar entrevista extendida",
        command=lambda: procesar_entrevista(),
        font=("Arial", 11, "bold"),
        bg="#FF9800", fg="white",
        activebackground="#F57C00",
        padx=15, pady=8
    ).pack(pady=8)

    # Botón cerrar
    tk.Button(
        win,
        text="❌ Cerrar",
        command=win.destroy,
        font=("Arial", 11, "bold"),
        bg="#f44336", fg="white",
        activebackground="#d32f2f",
        padx=15, pady=11
    ).pack(pady=20)

# -------------------------------------------------------------------------------
# NUEVO: Implementación de "Agregar actividad y evidencia"
# -------------------------------------------------------------------------------

def localizar_carpeta_final_alumno(ruta_carpeta_global, matricula, nombre_may, programa):
    """
    Busca la carpeta del alumno en la carpeta maestra y desciende
    automáticamente si hay solo una carpeta interna.
    """
    ruta_alumno = buscar_carpeta_alumno(ruta_carpeta_global, matricula, nombre_may, programa)

    if not os.path.isdir(ruta_alumno):
        raise FileNotFoundError(f"No se encontró la carpeta del alumno: {ruta_alumno}")

    # Listar subcarpetas
    subdirs = [d for d in os.listdir(ruta_alumno) if os.path.isdir(os.path.join(ruta_alumno, d))]
    
    # Si hay solo una subcarpeta, descendemos a ella
    if len(subdirs) == 1:
        ruta_alumno = os.path.join(ruta_alumno, subdirs[0])

    return ruta_alumno


def agregar_actividad_y_evidencia(datos):
    """Flujo completo para un único alumno: actualizar seguimiento en Word + añadir evidencia."""
    global seleccion_primera_vez
    global ruta_carpeta_global
    global ruta_excel_principal

    try:
        nombre = datos["Nombre completo"][0]
        nombre_may = nombre.upper()
        matricula = datos.get("Matrícula", [""])[0]
        programa = datos.get("Programa", [""])[0]

        if seleccion_primera_vez:

            # 1) Seleccionar carpeta maestra y Excel semanal
            ruta_carpeta_global = seleccionar_carpeta("Selecciona la carpeta MAESTRA que contiene todas las carpetas de alumnos:")
            #ruta_carpeta_global = ruta_carpeta_global
            ruta_excel_principal = seleccionar_excel_tabla_reporte() #Validar que si se haya seleccionado un excel
            seleccion_primera_vez = False

        # 2) Leer reporte y extraer fila más reciente del alumno
        df_rep = leer_tabla_reporte(ruta_excel_principal)
        fila = filtrar_fila_reciente_por_alumno(df_rep, nombre)

        # 3) Ubicar carpeta final del alumno y archivo de seguimiento
        ruta_alumno = localizar_carpeta_final_alumno(ruta_carpeta_global, matricula, nombre_may, programa)
        ruta_docx = encontrar_archivo_seguimiento(ruta_alumno)

        # 4) Insertar actividad en la tabla del Word
        insertar_actividad_en_word(ruta_docx, fila, nombre_may)

        # 5) Evidencias: ubicar carpeta y solicitar imagen + actividad
        carpeta_evid = encontrar_carpeta_evidencias(ruta_alumno)
        titulo = f"Selecciona la evidencia para {nombre}"
        ruta_img = filedialog.askopenfilename(title=titulo, filetypes=[
            ("Imágenes", "*.jpg;*.jpeg;*.png;*.heic;*.webp;*.bmp")
        ])
        if not ruta_img:
            messagebox.showinfo("Aviso", "No se seleccionó imagen de evidencia. Proceso de evidencia cancelado.")
        else:
            actividad = simpledialog.askstring("Actividad", "Nombre de la actividad para la evidencia:")
            if not actividad:
                messagebox.showwarning("Aviso", "No ingresaste el nombre de la actividad. Se canceló el guardado de evidencia.")
            else:
                iniciales = obtener_iniciales(nombre)
                ruta_img_dest = copiar_y_renombrar_evidencia(carpeta_evid, ruta_img, actividad, iniciales) ##¿?
                #messagebox.showinfo("Evidencia", f"Evidencia copiada y renombrada en:\n{destino}")

        messagebox.showinfo("Listo", "Actividad agregada al seguimiento y evidencia procesada correctamente.")

    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un problema: {e}")


# Placeholder: botón pendiente

def agregar_entrevista_extendida(datos):
    """Flujo completo para un único alumno: actualizar seguimiento en Word + añadir evidencia."""
    global seleccion_primera_vez
    global ruta_carpeta_global
    global ruta_excel_principal

    try:
        nombre = datos["Nombre completo"][0]
        nombre_may = nombre.upper()
        matricula = datos.get("Matrícula", [""])[0]
        programa = datos.get("Programa", [""])[0]

        if seleccion_primera_vez:

            # 1) Seleccionar carpeta maestra y Excel semanal
            ruta_carpeta_global = seleccionar_carpeta("Selecciona la carpeta MAESTRA que contiene todas las carpetas de alumnos:")
            #ruta_carpeta_global = ruta_carpeta_global
            ruta_excel_principal = seleccionar_excel_tabla_reporte() #Validar que si se haya seleccionado un excel
            seleccion_primera_vez = False

        # 1) Ubicar carpeta final del alumno y archivo de entrevista
        ruta_alumno = localizar_carpeta_final_alumno(ruta_carpeta_global, matricula, nombre_may, programa)
        ruta_docx = encontrar_archivo_entrevista(ruta_alumno)

        if not ruta_docx:
            messagebox.showerror("Error", f"No se encontró archivo de entrevista para {nombre}")
            return

        # 2) Seleccionar imágenes
        rutas_imagenes = filedialog.askopenfilenames(
            title=f"Selecciona las imágenes de la entrevista de {nombre}",
            filetypes=[("Archivos de imagen", "*.jpg *.jpeg *.png *.bmp *.tiff")]
        )
        if not rutas_imagenes:
            return

        # Ordenar imágenes por fecha de modificación (más antigua primero)
        rutas_imagenes = sorted(rutas_imagenes, key=os.path.getmtime)

        # 3) Insertar imágenes en el documento
        doc = Document(ruta_docx)
        section = doc.sections[0]

        # Dimensiones de página
        ancho_pagina = section.page_width.inches
        alto_pagina = section.page_height.inches
        margen_izq = section.left_margin.inches
        margen_der = section.right_margin.inches
        margen_sup = section.top_margin.inches
        margen_inf = section.bottom_margin.inches

        ancho_max = ancho_pagina - (margen_izq + margen_der)
        alto_max = alto_pagina - (margen_sup + margen_inf)

        for i, img_path in enumerate(rutas_imagenes):
            if i > 0:
                # Insertar salto de página limpio antes de la imagen
                doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
                espacio_disp = alto_max
            else:
                # Primera imagen: reservar espacio por texto inicial
                espacio_disp = alto_max - 2.0
                if espacio_disp < alto_max * 0.5:
                    espacio_disp = alto_max * 0.5

            # Abrir imagen y calcular escalado
            with Image.open(img_path) as im:
                w, h = im.size
                dpi = im.info.get("dpi", (96, 96))
                w_in = w / dpi[0]
                h_in = h / dpi[1]

                factor = min(ancho_max / w_in, espacio_disp / h_in)
                ancho_final = Inches(w_in * factor)
                alto_final = Inches(h_in * factor)

            # Insertar imagen
            doc.add_picture(img_path, width=ancho_final, height=alto_final)

        # 4) Guardar cambios
        doc.save(ruta_docx)
        messagebox.showinfo("Éxito", f"Entrevista actualizada con evidencias para {nombre}")

    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un problema al generar las entrevistas: {e}")

def previsualizar_datos_excel(datos_fila):
    """
    Muestra una ventana con los datos que se van a insertar en Word.
    `datos_fila` debe ser un diccionario con las claves: 'Fecha', 'Tema', 'Situación'.
    """
    ventana = tk.Toplevel()
    ventana.title("Previsualización de datos (DEBUG)")
    ventana.geometry("400x250")
    
    tk.Label(ventana, text="Datos a insertar en Word:", font=("Arial", 12, "bold")).pack(pady=10)

    tk.Label(ventana, text=f"Fecha de atención: {datos_fila['Fecha']}").pack(pady=5)
    tk.Label(ventana, text=f"Tema o asunto tratado: {datos_fila['Tema']}").pack(pady=5)
    tk.Label(ventana, text=f"Situación del alumno: {datos_fila['Situación']}").pack(pady=5)
    
    tk.Button(ventana, text="OK, continuar", command=ventana.destroy).pack(pady=15)
    ventana.grab_set()
    ventana.mainloop()

# -------------------------------------------------------------------------------
# Utilidades ya existentes para reemplazos y fechas
# -------------------------------------------------------------------------------

dias_semana = {"lunes": 0, "martes": 1, "miércoles": 2, "jueves": 3, "viernes": 4, "sábado": 5, "domingo": 6}


def proxima_reunion(dia, hora_inicio):
    if not dia or pd.isna(dia) or str(dia).lower() == 'nan':
        return "FECHA NO DEFINIDA"
    hoy = datetime.now()
    dia_actual = hoy.weekday()
    hora_act = hoy.time()
    partes = str(hora_inicio).split(":")
    try:
        h = int(partes[0])
        m = int(partes[1])
    except (ValueError, IndexError):
        h, m = 0, 0
    hora_ini = time(h, m)
    dia_reu = dias_semana.get(str(dia).lower().strip(), 0)
    if dia_reu < dia_actual or (dia_reu == dia_actual and hora_ini <= hora_act):
        dias_para = 7 - (dia_actual - dia_reu)
    else:
        dias_para = dia_reu - dia_actual
    prox = hoy + timedelta(days=dias_para)
    return f"{str(dia).upper()} {prox.day} DE {prox.strftime('%B').upper()} A LAS {hora_inicio}"


def formatear_hora(valor):
    if isinstance(valor, (datetime, time)):
        return valor.strftime("%H:%M")
    try:
        valor_str = str(valor)
        if ":" in valor_str:
            partes = valor_str.split(":")
            if len(partes) >= 2:
                return f"{partes[0].zfill(2)}:{partes[1].zfill(2)}"
        float_val = float(valor)
        frac_dia = float_val * 24
        horas = int(frac_dia)
        minutos = int((frac_dia - horas) * 60)
        return f"{horas:02d}:{minutos:02d}"
    except (ValueError, TypeError):
        return "00:00"


def reemplazar_en_docx(doc, markers):
    for para in doc.paragraphs:
        for mk, txt in markers.items():
            if mk in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if mk in inline[i].text:
                        text = inline[i].text.replace(mk, str(txt))
                        inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for mk, txt in markers.items():
                        if mk in para.text:
                            inline = para.runs
                            for i in range(len(inline)):
                                if mk in inline[i].text:
                                    text = inline[i].text.replace(mk, str(txt))
                                    inline[i].text = text


# -------------------------------------------------------------------------------
# App Tkinter
# -------------------------------------------------------------------------------

def ejecutar_y_ocultar(func, win):
    win.withdraw()
    func()


def iniciar_app():
    global root
    root = tk.Tk()
    root.title("Gestión de Expedientes")
    root.geometry("300x200")
    root.withdraw()
    ventana_bienvenida()
    root.mainloop()


def ventana_bienvenida():
    win = tk.Toplevel(root)
    win.title("Bienvenida")

    # Tamaño fijo
    ancho, alto = 400, 250
    # Calculamos el centro de la pantalla
    x = (win.winfo_screenwidth() // 2) - (ancho // 2)
    y = (win.winfo_screenheight() // 2) - (alto // 2)
    win.geometry(f"{ancho}x{alto}+{x}+{y}")
    win.resizable(False, False)

    # Mensaje de bienvenida
    tk.Label(
        win, 
        text="Bienvenido al gestor de expedientes", 
        font=("Arial", 14, "bold"),
        pady=20
    ).pack()

    # Botón destacado
    boton = tk.Button(
        win, 
        text="🚀 Iniciar", 
        command=lambda: ejecutar_y_ocultar(cargar_datos_excel, win),
        font=("Arial", 12, "bold"),
        bg="#4CAF50",      # verde moderno
        fg="white",
        activebackground="#45a049",
        padx=20, pady=10
    )
    boton.pack(pady=30)


if __name__ == "__main__":
    iniciar_app()
