#!/usr/bin/env python3
import os
import sys
import platform
import shutil
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
import pickle
import json
import locale
from datetime import datetime, timedelta, time
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from PIL import Image
from docx.shared import Inches
from docx import Document

# Variable global para reutilizar la ruta
ruta_excel_principal = None
ruta_carpeta_global = None
ruta_plantillas = None
seleccion_primera_vez = True

# DEBUG flag (poner True solo mientras se depura)
DEBUG_LEER_REPORTE = True

# -------------------------------------------------------------------------------
# CONFIGURACIÓN Y ESTILOS UI
# -------------------------------------------------------------------------------

# Lista centralizada de tutores
LISTA_TUTORES = [
    "Diana KKufnedkjfsl",
    "Luis Polero haytrg",
    "Yenifer frlokmve",
    "Alan vedlkmlmñ",
    "Veda wdscaasfv gfbgf"
]

# Colores del tema oscuro
COLOR_BG = "#2b2b2b"
COLOR_FG = "#ffffff"
COLOR_ACCENT = "#4CAF50"  # Verde moderno
COLOR_ACCENT_HOVER = "#45a049"
COLOR_SECONDARY = "#2196F3" # Azul
COLOR_WARNING = "#FF9800" # Naranja
COLOR_ERROR = "#f44336" # Rojo
COLOR_ENTRY_BG = "#3d3d3d"

# Detección de Sistema y Fuentes
SYSTEM_OS = platform.system()
if SYSTEM_OS == "Windows":
    FONT_FAMILY = "Segoe UI"
elif SYSTEM_OS == "Darwin": # macOS
    FONT_FAMILY = "Helvetica"
else:
    FONT_FAMILY = "Arial"

FONT_MAIN = (FONT_FAMILY, 10)
FONT_HEADER = (FONT_FAMILY, 12, "bold")
FONT_TITLE = (FONT_FAMILY, 14, "bold")

def aplicar_estilo(root):
    """Aplica el tema oscuro y estilos personalizados a la aplicación."""
    style = ttk.Style(root)
    style.theme_use("clam") # Base theme that supports color customization well

    # Configuración general
    style.configure(".", background=COLOR_BG, foreground=COLOR_FG, font=FONT_MAIN)
    
    # TFrames
    style.configure("TFrame", background=COLOR_BG)
    
    # TLabels
    style.configure("TLabel", background=COLOR_BG, foreground=COLOR_FG, font=FONT_MAIN)
    style.configure("Header.TLabel", font=FONT_HEADER)
    style.configure("Title.TLabel", font=FONT_TITLE)

    # TButtons
    style.configure("TButton", 
                    background=COLOR_ENTRY_BG, 
                    foreground=COLOR_FG, 
                    borderwidth=0, 
                    focuscolor=COLOR_ACCENT,
                    font=FONT_MAIN,
                    padding=6)
    style.map("TButton", 
              background=[("active", "#505050"), ("pressed", "#606060")])

    # Accent Button (Green)
    style.configure("Accent.TButton", 
                    background=COLOR_ACCENT, 
                    foreground="white",
                    font=(FONT_FAMILY, 10, "bold"))
    style.map("Accent.TButton", 
              background=[("active", COLOR_ACCENT_HOVER)])

    # Secondary Button (Blue)
    style.configure("Secondary.TButton", 
                    background=COLOR_SECONDARY, 
                    foreground="white",
                    font=(FONT_FAMILY, 10, "bold"))
    style.map("Secondary.TButton", 
              background=[("active", "#1976D2")])
    
    # Danger Button (Red)
    style.configure("Danger.TButton", 
                    background=COLOR_ERROR, 
                    foreground="white",
                    font=(FONT_FAMILY, 10, "bold"))
    style.map("Danger.TButton", 
              background=[("active", "#d32f2f")])

    # Scrollbar
    style.configure("Vertical.TScrollbar", 
                    background=COLOR_ENTRY_BG, 
                    troughcolor=COLOR_BG,
                    arrowcolor=COLOR_FG)

    # Checkbutton
    style.configure("TCheckbutton", background=COLOR_BG, foreground=COLOR_FG, font=FONT_MAIN)
    style.map("TCheckbutton", background=[("active", COLOR_BG)])

    # Configurar root
    root.configure(bg=COLOR_BG)

def centrar_ventana(win, ancho, alto):
    """Centra una ventana en la pantalla."""
    x = (win.winfo_screenwidth() // 2) - (ancho // 2)
    y = (win.winfo_screenheight() // 2) - (alto // 2)
    win.geometry(f"{ancho}x{alto}+{x}+{y}")

def force_focus(window):
    """Fuerza el foco en la ventana, especialmente para macOS."""
    window.lift()
    window.focus_force()
    if SYSTEM_OS == "Darwin":
        # Truco para traer al frente en macOS
        window.attributes('-topmost', True)
        window.after_idle(window.attributes, '-topmost', False)
        # Asegurar que se levante sobre otras apps (opcional, usar con cuidado)
        try:
            os.system('''/usr/bin/osascript -e 'tell app "Finder" to set frontmost of process "Python" to true' ''')
        except:
            pass

def crear_ventana_toplevel(titulo, ancho=400, alto=300):
    """Crea una ventana Toplevel con el estilo configurado."""
    win = tk.Toplevel(root)
    win.title(titulo)
    win.configure(bg=COLOR_BG)
    centrar_ventana(win, ancho, alto)
    force_focus(win)
    return win

# -------------------------------------------------------------------------------
# Utilidades generales
# -------------------------------------------------------------------------------

def obtener_iniciales(nombre_completo: str) -> str:
    partes = str(nombre_completo).split()
    return ''.join([p[0].upper() for p in partes if p])

try:
    locale.setlocale(locale.LC_TIME, "es_ES.UTF-8")
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, "Spanish_Spain.1252")
    except locale.Error:
        pass

CONFIG_FILE = "config.json"


def guardar_config(carpeta_raiz):
    with open(CONFIG_FILE, "w") as f:
        json.dump({"carpeta_raiz": carpeta_raiz}, f)


def cargar_config():
    try:
        with open(CONFIG_FILE, "r") as f:
            return json.load(f).get("carpeta_raiz", "")
    except FileNotFoundError:
        return ""


# -------------------------------------------------------------------------------
# Utilidades específicas para agregar actividad y evidencia
# -------------------------------------------------------------------------------

def _map_nombre_columna(raw):
    """Mapea nombres de columna detectados a los nombres esperados de manera tolerante."""
    if raw is None:
        return raw
    r = str(raw).strip().lower()
    # coincidencias por palabra clave
    if "nombre" in r and "apellidos" not in r and "apellido" not in r:
        return "Nombre"
    if "apellido" in r:
        return "Apellidos"
    if "fecha" in r and "atenc" in r:
        return "Fecha de atención"
    if "fecha" in r and "atenc" not in r:
        # aceptamos 'Fecha' como 'Fecha de atención'
        return "Fecha de atención"
    if "tema" in r or "asunto" in r:
        return "Tema o asunto tratado"
    if "situac" in r or "situa" in r:
        return "Situación del alumno"
    # por defecto: devolver título "limpio" (capitalizado)
    return str(raw).strip()

def normalizar_columnas(df: pd.DataFrame) -> pd.DataFrame:
    """Limpia espacios en los encabezados y estandariza nombres esperados."""
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    # Aliases tolerantes
    renombres = {
        "Tema o asunto tratado": "Tema o asunto tratado",
        "Tema o asunto tratado ": "Tema o asunto tratado",
        "Situación del alumno": "Situación del alumno",
        "Fecha de atención": "Fecha de atención",
        "Nombre": "Nombre",
        "Apellidos": "Apellidos",
    }
    df.rename(columns=renombres, inplace=True)
    return df


def seleccionar_excel_tabla_reporte() -> str:
    global ruta_excel_principal
    if not ruta_excel_principal:
        raise SystemExit("No se seleccionó Excel semanal.")
    return ruta_excel_principal


def leer_tabla_reporte(ruta_excel_principal: str, hoja=2, DEBUG_LEER_REPORTE=False) -> pd.DataFrame:
    """
    Lee una hoja específica del Excel (por índice o nombre) y devuelve un DataFrame filtrado.
    Detecta automáticamente el encabezado en B11:J11 y valida que existan las columnas 'Nombre' y 'Apellidos'.
    """
    try:
        # Leer la hoja completa
        # Si 'hoja' es int, usa sheet_name por índice, si es str usa por nombre
        df_full = pd.read_excel(ruta_excel_principal, sheet_name=hoja, header=None, keep_default_na=False)
    except Exception as e:
        messagebox.showerror("Error de Lectura", f"No se pudo leer la hoja 3 del Excel.\n\nAsegúrate de que el archivo no esté corrupto.\n\nDetalle: {e}")
        raise

    # Tomar solo las columnas B:J y filas desde la fila 11 (índice 10) para detectar encabezado
    df_header = df_full.iloc[10, 1:11]  # B11:K11
    df_data = df_full.iloc[11:, 1:11]   # Datos debajo del encabezado
    df_data.columns = [str(c).strip() for c in df_header]

    # --- Validación de columnas críticas ---
    cols_lower = [c.lower() for c in df_data.columns]
    tiene_nombre = any(c == "nombre" for c in cols_lower)
    tiene_apellidos = any(c == "apellidos" or c == "apellido" for c in cols_lower)
    if not (tiene_nombre and tiene_apellidos):
        cols_found = ", ".join(df_data.columns.tolist())
        raise ValueError(
            f"La hoja 3 no contiene las columnas esperadas 'Nombre' y 'Apellidos'.\n"
            f"Columnas detectadas: {cols_found}"
        )

    # Generar columna de nombre completo (Reporte)
    df_data["Nombre completo (Reporte)"] = (
        df_data["Nombre"].astype(str).str.strip() + " " +
        df_data["Apellidos"].astype(str).str.strip()
    ).str.strip()

    return df_data.reset_index(drop=True)


def filtrar_fila_reciente_por_alumno(df_reporte: pd.DataFrame, nombre_objetivo: str) -> pd.Series:
    """
    Selecciona la fila más reciente del alumno en df_reporte comparando
    'Nombre completo (Reporte)'. Si no encuentra, lanza error con pistas.
    """
    if "Nombre completo (Reporte)" not in df_reporte.columns:
        # intentar reconstruir desde columnas detectables
        cols = df_reporte.columns.tolist()
        if any("nombre" in c.lower() for c in cols) and any("apell" in c.lower() for c in cols):
            ncol = next(c for c in cols if "nombre" in c.lower())
            acol = next(c for c in cols if "apell" in c.lower())
            df_reporte["Nombre completo (Reporte)"] = (df_reporte[ncol].fillna("").astype(str).str.strip() + " " + df_reporte[acol].fillna("").astype(str).str.strip()).str.strip()
        else:
            raise ValueError("No se pudo construir 'Nombre completo (Reporte)' desde las columnas existentes.")

    # Comparación insensible a mayúsculas y espacios
    objetivo_norm = str(nombre_objetivo).strip().upper()
    df_reporte["___NOM_SEARCH__"] = df_reporte["Nombre completo (Reporte)"].astype(str).str.strip().str.upper()

    df_al = df_reporte[df_reporte["___NOM_SEARCH__"] == objetivo_norm].copy()
    if df_al.empty:
        # Preparar mensaje útil para depurar: listar algunos nombres detectados
        uniques = df_reporte["Nombre completo (Reporte)"].dropna().unique().tolist()[:30]
        raise ValueError(f"No se encontraron registros para '{nombre_objetivo}' en TablaReporte.\n\nVerifica que el nombre en el Excel coincida exactamente.")
    # ordenar por Fecha de atención si existe
    if "Fecha de atención" in df_al.columns:
        df_al["Fecha de atención"] = pd.to_datetime(df_al["Fecha de atención"], errors="coerce")
        df_al.sort_values(by="Fecha de atención", inplace=True, na_position="first")
        fila = df_al.iloc[-1]
    else:
        fila = df_al.iloc[-1]
    # limpiar columna temporal
    df_reporte.drop(columns=["___NOM_SEARCH__"], errors="ignore", inplace=True)
    return fila


def buscar_carpeta_alumno(ruta_carpeta_global: str, matricula: str, nombre_mayus: str, programa: str) -> str:
    """
    Busca la carpeta del alumno en la ruta global usando matrícula, nombre y programa.
    Retorna la ruta completa si existe, o lanza FileNotFoundError.
    """
    base_folder = f"{matricula}_{nombre_mayus}_{programa}"
    ruta = os.path.join(ruta_carpeta_global, base_folder)
    if not os.path.isdir(ruta):
        raise FileNotFoundError(f"No se encontró la carpeta del alumno: {base_folder}")
    return ruta


def encontrar_archivo_seguimiento(ruta_alumno: str) -> str:
    """Busca el archivo Word que comienza con '5_Seguimiento individual_' en el árbol del alumno."""
    for root_dir, _, files in os.walk(ruta_alumno):
        for f in files:
            if f.lower().endswith(".docx") and f.startswith("5_Seguimiento individual_"):
                return os.path.join(root_dir, f)
    raise FileNotFoundError("No se encontró un archivo que comience con '5_Seguimiento individual_'.")

def encontrar_archivo_entrevista(ruta_alumno: str) -> str:
    """Busca el archivo Word que comienza con '3_Entrevista extendida_' en el árbol del alumno."""
    for root_dir, _, files in os.walk(ruta_alumno):
        for f in files:
            if f.lower().endswith(".docx") and f.startswith("3_Entrevista extendida_"):
                return os.path.join(root_dir, f)
    raise FileNotFoundError("No se encontró un archivo que comience con '3_Entrevista extendida_'.")


def encontrar_carpeta_evidencias(ruta_alumno: str) -> str:
    """
    Busca la carpeta que comienza con '6_Evidencias_' dentro de la carpeta del alumno.
    """
    for root_dir, dirs, _ in os.walk(ruta_alumno):
        for d in dirs:
            if d.startswith("6_Evidencias_"):
                return os.path.join(root_dir, d)
    raise FileNotFoundError("No se encontró carpeta que comience con '6_Evidencias_'.")

def crear_imagen_blanca(ruta_destino, ancho=200, alto=200):
    """Genera una imagen PNG en blanco en la ruta indicada."""
    img = Image.new("RGB", (ancho, alto), (255, 255, 255))
    img.save(ruta_destino)
    return ruta_destino

def primera_fila_col4_vacia(tabla: docx.table.Table) -> int:
    """Devuelve el índice de fila (0-based) donde la columna 4 esté vacía por primera vez.
    Si no hay, devuelve len(rows) para que podamos anexar una nueva fila."""
    for i, row in enumerate(tabla.rows):
        celdas = row.cells
        if len(celdas) >= 4:
            texto_col4 = celdas[3].text.strip()
            if texto_col4 == "":
                return i
    return len(tabla.rows)


def asegurar_filas(tabla: docx.table.Table, filas_necesarias: int):
    """Agrega filas a la tabla hasta alcanzar el total filas_necesarias (si hace falta)."""
    while len(tabla.rows) < filas_necesarias:
        tabla.add_row()

# -------------------------------------------------------------------------------
# GUI: Selección de Tutor y Asistencia
# -------------------------------------------------------------------------------

def seleccionar_tutor_gui(titulo="Seleccionar Tutor"):
    """
    Muestra una ventana moderna para seleccionar un tutor de la lista centralizada.
    Retorna el nombre seleccionado o None si se cancela.
    """
    win = crear_ventana_toplevel(titulo, ancho=400, alto=400)
    
    seleccion = tk.StringVar(value="")

    ttk.Label(win, text="Selecciona un Tutor:", style="Header.TLabel").pack(pady=15)

    frame_lista = ttk.Frame(win)
    frame_lista.pack(fill="both", expand=True, padx=20, pady=5)

    # Scrollbar para la lista si crece mucho
    canvas = tk.Canvas(frame_lista, bg=COLOR_BG, highlightthickness=0)
    scrollbar = ttk.Scrollbar(frame_lista, orient="vertical", command=canvas.yview, style="Vertical.TScrollbar")
    scrollable_frame = ttk.Frame(canvas)

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", width=360)
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")
    
    # Botones de tutores
    botones = []
    
    def marcar(nombre, btn_obj):
        seleccion.set(nombre)
        # Reset estilos
        for b in botones:
            b.configure(style="TButton")
        # Marcar activo
        btn_obj.configure(style="Accent.TButton")

    for tutor in LISTA_TUTORES:
        btn = ttk.Button(scrollable_frame, text=tutor)
        # Usamos closure para capturar el botón actual
        btn.configure(command=lambda t=tutor, b=btn: marcar(t, b))
        btn.pack(fill="x", pady=2)
        botones.append(btn)

    def confirmar():
        if not seleccion.get():
            messagebox.showwarning("Atención", "Por favor selecciona un tutor.")
            return
        win.destroy()

    ttk.Button(win, text="Confirmar Selección", style="Accent.TButton", command=confirmar).pack(pady=20)

    win.grab_set()
    win.wait_window()
    
    return seleccion.get() if seleccion.get() else None


def seleccionar_tutor_y_asistencia(nombre):
    """
    Muestra ventana para seleccionar tutor y si asistió el alumno.
    Retorna:
        nombre_tutor (str)
        asistio (bool)
    """
    ventana = crear_ventana_toplevel(f"Información para {nombre}", ancho=800, alto=450)
    
    nombre_tutor = tk.StringVar(value="")
    asistio = tk.StringVar(value="")

    # --- Sección Tutor ---
    ttk.Label(ventana, text=f"Tutor para {nombre}:", style="Header.TLabel").pack(pady=10)
    
    frame_tutores = ttk.Frame(ventana)
    frame_tutores.pack(pady=5, fill="x", padx=20)
    
    # Grid de tutores para mejor distribución
    botones_tutores = {}
    
    def marcar_tutor(tutor):
        nombre_tutor.set(tutor)
        for t, btn in botones_tutores.items():
            btn.configure(style="TButton")
        botones_tutores[tutor].configure(style="Accent.TButton")

    # Crear grid 2 columnas
    for i, t in enumerate(LISTA_TUTORES):
        btn = ttk.Button(frame_tutores, text=t, command=lambda t=t: marcar_tutor(t))
        btn.grid(row=i//2, column=i%2, padx=5, pady=5, sticky="ew")
        botones_tutores[t] = btn
    
    frame_tutores.columnconfigure(0, weight=1)
    frame_tutores.columnconfigure(1, weight=1)

    # --- Sección Asistencia ---
    ttk.Label(ventana, text="¿Asistió el alumno?", style="Header.TLabel").pack(pady=(20,5))
    frame_asistencia = ttk.Frame(ventana)
    frame_asistencia.pack(pady=5)

    botones_asistencia = {}

    def marcar_asistencia(valor):
        asistio.set(valor)
        # Reset visual
        for b in botones_asistencia.values():
            b.configure(style="TButton")
        
        # Estilo según selección
        estilo = "Accent.TButton" if valor == "Sí" else "Danger.TButton"
        botones_asistencia[valor].configure(style=estilo)

        # Auto-cerrar si ya hay tutor seleccionado? No, mejor esperar confirmación explícita o cerrar al final.
        # En el diseño anterior se cerraba al seleccionar asistencia. Mantengamos eso pero validando tutor.
        if not nombre_tutor.get():
            messagebox.showwarning("Falta Tutor", "Primero selecciona un tutor.")
            # Reset asistencia visualmente para no confundir
            asistio.set("") 
            botones_asistencia[valor].configure(style="TButton")
            return
        
        ventana.destroy()

    btn_si = ttk.Button(frame_asistencia, text="Sí", command=lambda: marcar_asistencia("Sí"))
    btn_si.pack(side="left", padx=5)
    botones_asistencia["Sí"] = btn_si

    btn_no = ttk.Button(frame_asistencia, text="No", command=lambda: marcar_asistencia("No"))
    btn_no.pack(side="left", padx=5)
    botones_asistencia["No"] = btn_no

    btn_na = ttk.Button(frame_asistencia, text="N/A", command=lambda: marcar_asistencia("N/A"))
    btn_na.pack(side="left", padx=5)
    botones_asistencia["N/A"] = btn_na

    btn_tarea = ttk.Button(frame_asistencia, text="Tarea", command=lambda: marcar_asistencia("Tarea"))
    btn_tarea.pack(side="left", padx=5)
    botones_asistencia["Tarea"] = btn_tarea

    btn_just = ttk.Button(frame_asistencia, text="Falta justificada", command=lambda: marcar_asistencia("Falta justificada"))
    btn_just.pack(side="left", padx=5)
    botones_asistencia["Falta justificada"] = btn_just

    ventana.grab_set()
    ventana.wait_window()

    if not nombre_tutor.get() or not asistio.get():
        # Si cerró la ventana sin seleccionar
        if messagebox.askyesno("Cancelar", "¿Deseas cancelar la operación para este alumno?"):
            raise InterruptedError("Operación cancelada por el usuario.")
        else:
            return seleccionar_tutor_y_asistencia(nombre) # Reintentar

    return nombre_tutor.get(), asistio.get()


def _escribir_renglon_seguimiento(tabla: docx.table.Table, fila: pd.Series, nombre_tutor, asistio):
    """
    Función interna para escribir una fila en la tabla Word sin abrir/guardar el archivo.
    Retorna True si insertó, False si detectó duplicado (opcional, por ahora inserta siempre).
    """
    col_fecha = 1
    col_tema = 2
    col_situacion = 3
    col_Asistencia = 4
    col_Tutor = 5

    # Convertir valores a string seguro
    fecha = fila.get("Fecha de atención")
    Hora = str(fila.get("Hora"))

    if pd.notna(fecha):
        if isinstance(fecha, pd.Timestamp):
            fecha = fecha.strftime("%d/%m/%Y") + " " + Hora
        else:
            fecha = str(fecha).strip()
    else:
        fecha = ""

    # Fix: Handle 'nan' string or NaN values explicitly
    def clean_text(val):
        s = str(val).strip()
        if s.lower() == 'nan':
            return ""
        return s

    tema = clean_text(fila.get("Tema o asunto tratado", ""))
    situacion = clean_text(fila.get("Situación del alumno", ""))

    # --- DETECCION DE DUPLICADOS ---
    # Si ya existe EXACTAMENTE la misma fecha y tema, ignoramos.
    for r in tabla.rows:
        try:
            cf = r.cells[col_fecha].text.strip()
            ct = r.cells[col_tema].text.strip()
            if cf == fecha and ct == tema:
                 # Ya existe
                 print(f"  [Info] Registro duplicado omitido: {fecha} - {tema}")
                 return False
        except IndexError:
            pass

    # Buscar índice de la primera fila vacía en la columna 2 (Tema)
    idx_insert = None
    for i, row in enumerate(tabla.rows):
        celda = row.cells[col_tema]
        if not celda.text.strip():
            idx_insert = i
            break
    
    if idx_insert is not None:
        if asistio == "N/A":
            # Si es N/A, insertar nueva fila antes de la vacía para desplazarla
            fila_existente = tabla.rows[idx_insert]
            fila_insert = tabla.add_row()
            fila_existente._tr.addprevious(fila_insert._tr)
        else:
            # Si es Sí/No, usar la fila vacía existente
            fila_insert = tabla.rows[idx_insert]
    else:
        # Si no hay vacías, agregar al final
        fila_insert = tabla.add_row()

    fila_insert.cells[col_fecha].text = fecha
    fila_insert.cells[col_tema].text = tema
    fila_insert.cells[col_situacion].text = situacion
    fila_insert.cells[col_Asistencia].text = str(asistio)
    fila_insert.cells[col_Tutor].text = str(nombre_tutor)

    # Estilo
    for cell in [fila_insert.cells[col_fecha], fila_insert.cells[col_tema],
                fila_insert.cells[col_situacion], fila_insert.cells[col_Asistencia],
                fila_insert.cells[col_Tutor]]:
        
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
    
    return True


def insertar_actividad_en_word(ruta_docx: str, fila: pd.Series, nombre, tutor_predefinido=None, asistencia_predefinida=None):
    """
    Inserta los datos de una fila de reporte en la tabla del Word.
    Si se dan tutor_predefinido y asistencia_predefinida, NO abre ventana de diálogo (modo silencioso).
    Retorna: tuple: (nombre_tutor, asistio)
    """
    if not os.path.isfile(ruta_docx):
        raise FileNotFoundError(f"No se encontró el archivo Word: {ruta_docx}")

    # Solicitar tutor y asistencia o usar predefinidos
    if tutor_predefinido is not None and asistencia_predefinida is not None:
        nombre_tutor = tutor_predefinido
        asistio = asistencia_predefinida
    else:
        try:
            nombre_tutor, asistio = seleccionar_tutor_y_asistencia(nombre)
        except InterruptedError:
            return None, None # Señal de cancelación

    try:
        doc = docx.Document(ruta_docx)
    except Exception as e:
        messagebox.showerror("Error de Archivo", f"No se pudo abrir el documento Word.\n\nCierra el archivo si lo tienes abierto.\n\nDetalle: {e}")
        raise RuntimeError(f"No se pudo abrir el documento Word: {e}")

    # Buscar la primera tabla
    if not doc.tables:
        raise ValueError(f"El documento {ruta_docx} no contiene tablas.")

    tabla = doc.tables[0]
    
    # Usar lógica extraída
    # Si es Tarea o Falta justificada, guardar como "Sí" en el documento
    asistio_para_reporte = asistio
    if asistio in ["Tarea", "Falta justificada"]:
        asistio_para_reporte = "Sí"

    _escribir_renglon_seguimiento(tabla, fila, nombre_tutor, asistio_para_reporte)
    
    try:
        doc.save(ruta_docx)
    except PermissionError:
        messagebox.showerror("Error de Permiso", f"No se pudo guardar el archivo Word.\n\n¡CIERRA EL ARCHIVO WORD SI LO TIENES ABIERTO!\n\nRuta: {ruta_docx}")
        raise

    return nombre_tutor, asistio


def proximo_numero_sesion(carpeta_evidencias: str, iniciales: str) -> int:
    """Calcula el siguiente número de sesión buscando archivos existentes 'sesión_#_..._INICIALES'."""
    extensiones = (".jpg", ".jpeg", ".png")
    archivos = [f for f in os.listdir(carpeta_evidencias) if f.lower().endswith(extensiones)]
    return len(archivos) + 1


def copiar_y_renombrar_evidencia(carpeta_evidencias: str, ruta_imagen_src: str, actividad: str, iniciales: str) -> str:
    num_sesion = proximo_numero_sesion(carpeta_evidencias, iniciales)
    actividad_limpia = "_".join(str(actividad).strip().split())
    nombre_dest = f"Sesión{num_sesion}_{actividad_limpia}_{iniciales}{os.path.splitext(ruta_imagen_src)[1].lower()}"
    ruta_dest = os.path.join(carpeta_evidencias, nombre_dest)
    shutil.copy2(ruta_imagen_src, ruta_dest)
    return ruta_dest


# -------------------------------------------------------------------------------
# Módulo crear_expedientes
# -------------------------------------------------------------------------------

def crear_expedientes(datos):
    """
    Crea la estructura de carpetas y copia plantillas para los alumnos seleccionados.
    """
    global seleccion_primera_vez
    global ruta_carpeta_global
    global ruta_plantillas

    if seleccion_primera_vez:
        ruta_carpeta_global = seleccionar_carpeta_raiz("Selecciona la carpeta donde se crearán los expedientes:")
        plantilla_dir = seleccionar_carpeta("Selecciona la carpeta donde están las plantillas:")
        ruta_plantillas = plantilla_dir
        seleccion_primera_vez = False

    prefijo = simpledialog.askstring("Prefijo de Carpeta", "Introduce el prefijo antes de '@' en la carpeta interna:")

    if not prefijo:
        messagebox.showerror("Error", "No ingresaste un prefijo. Proceso cancelado.")
        return

    for i in range(len(datos.get("Nombre completo", []))):
        nombre = datos["Nombre completo"][i].upper()
        matricula = datos.get("Matrícula", [None])[i]
        programa = datos.get("Programa", [None])[i]
        periodo = datos.get("Período", [None])[i]

        if not (nombre and matricula and programa and periodo):
            continue

        base_folder = f"{matricula}_{nombre}_{programa}"
        ruta_alumno = os.path.join(ruta_carpeta_global, base_folder)
        os.makedirs(ruta_alumno, exist_ok=True)

        iniciales = obtener_iniciales(nombre)
        interna = f"{prefijo}@{periodo}_{matricula}_{iniciales}_{programa}"
        ruta_interna = os.path.join(ruta_alumno, interna)
        os.makedirs(ruta_interna, exist_ok=True)

        # Copiar/renombrar plantillas y carpetas
        for item in os.listdir(plantilla_dir):
            src_path = os.path.join(plantilla_dir, item)
            if os.path.isfile(src_path):
                nuevo_nombre = f"{os.path.splitext(item)[0]}_{iniciales}_{programa}_{periodo}{os.path.splitext(item)[1]}"
                dest_path = os.path.join(ruta_interna, nuevo_nombre)
                shutil.copy2(src_path, dest_path)
            elif os.path.isdir(src_path):
                nuevo_nombre_carpeta = f"{item}_{iniciales}_{programa}_{periodo}"
                dest_dir_path = os.path.join(ruta_interna, nuevo_nombre_carpeta)
                os.makedirs(dest_dir_path, exist_ok=True)
                for archivo_interno in os.listdir(src_path):
                    src_file_path = os.path.join(src_path, archivo_interno)
                    if archivo_interno.lower().endswith((".jpeg", ".jpg")):
                        nombre_base, extension = os.path.splitext(archivo_interno)
                        nuevo_nombre_archivo = f"{nombre_base}_{iniciales}{extension}"
                    else:
                        nuevo_nombre_archivo = archivo_interno
                    dest_file_path = os.path.join(dest_dir_path, nuevo_nombre_archivo)
                    shutil.copy2(src_file_path, dest_file_path)

    messagebox.showinfo("Proceso Completo", "Los expedientes fueron creados exitosamente.")
    mostrar_crear_opciones(datos)


def mostrar_crear_opciones(datos):
    win = crear_ventana_toplevel("Opciones después de crear expedientes", ancho=350, alto=250)

    ttk.Label(win, text="¿Qué deseas hacer ahora?", style="Header.TLabel").pack(pady=20)
    
    ttk.Button(
        win,
        text="Llenar Documentos",
        style="Accent.TButton",
        command=lambda: (llenar_documentos(datos, win), win.destroy())
    ).pack(pady=10, fill="x", padx=40)
    
    ttk.Button(win, text="Continuar sin llenar", command=win.destroy).pack(pady=10, fill="x", padx=40)


# -------------------------------------------------------------------------------
# Módulo llenar_documentos
# -------------------------------------------------------------------------------

def llenar_documentos(datos, ventana=None):
    """
    Rellena los documentos Word en las carpetas de los alumnos con los datos del Excel.
    """
    raiz = cargar_config()
    if not raiz:
        messagebox.showerror("Error", "No se encontró carpeta raíz. Ejecuta primero crear_expedientes.")
        return

    # NUEVO: Usar selector GUI en lugar de input manual
    NombreTutor = seleccionar_tutor_gui("Selecciona el Tutor para los documentos")
    if not NombreTutor:
        messagebox.showwarning("Cancelado", "No se seleccionó tutor. Se cancela el llenado.")
        return

    for i in range(len(datos.get("Nombre completo", []))):
        nombre = datos["Nombre completo"][i]
        matricula = datos["Matrícula"][i]
        programa = datos["Programa"][i]
        periodo = datos["Período"][i]
        dia = datos.get("Dia", [None]*len(datos["Nombre completo"]))[i]
        hora_ini_raw = datos.get("Hora inicio", [None]*len(datos["Nombre completo"]))[i]
        horafin_raw = datos.get("Hora final", [None]*len(datos["Nombre completo"]))[i]

        hora_ini = formatear_hora(hora_ini_raw)
        horafin = formatear_hora(horafin_raw)

        horario = f"{dia} {hora_ini}-{horafin}"
        area = datos.get("Área de intervención", [None])[i] or ""
        edad = datos.get("Edad", [None])[i] or ""
        semestre = datos.get("Semestre", [None])[i] or ""
        habilidades = datos.get("Habilidades", [None])[i] or ""
        grupoProv = datos.get("Grupo", [None])[i] or ""
        titularDir = datos.get("Titular/Director", [None])[i] or ""
        situacion = datos.get("Situación", [None])[i] or ""
        promedio = datos.get("Promedio Anterior", [None])[i] or ""

        base_folder = f"{matricula}_{nombre.upper()}_{programa}"
        ruta_base = os.path.join(raiz, base_folder)

        programa_bach = ""

        if programa == "BUNLA" or programa == "BIUNLA" or programa == "BUNLAV":
            programa_bach = programa
            programa = ""
        else:
            grupoProv = ""

        if not os.path.isdir(ruta_base):
            print(f"Advertencia: No se encontró la carpeta para {nombre}. Saltando...")
            continue

        for root_dir, _, files in os.walk(ruta_base):
            for file in files:
                if not file.lower().endswith(".docx"):
                    continue

                ruta_doc = os.path.join(root_dir, file)
                try:
                    doc = docx.Document(ruta_doc)

                    markers = {
                        "{NOMBRE}": nombre,
                        "{NOMBREMAY}": nombre.upper(),
                        "{MATRICULA}": matricula,
                        "{PROGRAMA}": programa,
                        "{PROGRAMABACH}": programa_bach,
                        "{GRUPO}": grupoProv,
                        "{TitularDir}": titularDir,
                        "{SITUACIÓN}": situacion,
                        "{PERIODO}": periodo,
                        "{PROMEDIO}": promedio,
                        "{EDAD}": str(edad)+" años",
                        "{AREA}": area,
                        "{TUTOR}": NombreTutor,
                        "{SEMESTRE}": str(semestre),
                        "{HORARIO}": horario,
                        "{HABILIDADES}": habilidades,
                        "{FECHA}": datetime.now().strftime("%d/%m/%Y"),
                        "{FECHALARGA}": datetime.now().strftime("%d de %B de %Y"),
                        "{FECHADIASNANO}": datetime.now().strftime("%A %d de %B").upper(),
                        "{HORAINICIO}": hora_ini,
                        "{HORAFINAL}": horafin,
                        "{DIAINICIO}": dia,
                        "{PROXREUNION}": proxima_reunion(dia, hora_ini)
                    }

                    reemplazar_en_docx(doc, markers)
                    doc.save(ruta_doc)
                except Exception as e:
                    print(f"No se pudo procesar el archivo {ruta_doc}. Error: {e}")

    messagebox.showinfo("Proceso Completo", "Los documentos fueron llenados exitosamente.")


# -------------------------------------------------------------------------------
# Resto de utilidades existentes
# -------------------------------------------------------------------------------

def enviar_bienvenida(datos):
    print("[Correo] Bienvenida:", datos)


def enviar_extranamiento(datos):
    print("[Correo] Extranamiento:", datos)


def añadir_a_pit(datos):
    pass  # TODO: implementar lógica PIT


def imprimir_entrevista_b1(): print("Imprimiendo Entrevista B1...")

def imprimir_entrevista_bcont(): print("Imprimiendo Entrevista BCont...")

def imprimir_entrevista_l1(): print("Imprimiendo Entrevista L1...")

def imprimir_entrevista_lcont(): print("Imprimiendo Entrevista LCont...")

def imprimir_reglas(): print("Imprimiendo Reglas...")


def seleccionar_carpeta(mensaje):
    messagebox.showinfo("Instrucción", mensaje)
    ruta = filedialog.askdirectory(title=mensaje)
    if not ruta:
        sys.exit(0)
    return ruta


def seleccionar_carpeta_raiz(mensaje):
    messagebox.showinfo("Instrucción", mensaje)
    ruta = filedialog.askdirectory(title=mensaje)
    if not ruta:
        sys.exit(0)
    guardar_config(ruta)
    return ruta


def cargar_datos_excel():
    global ruta_excel_principal
    messagebox.showinfo("Instrucción", "Selecciona el archivo Excel con datos de los alumnos")
    ruta = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx;*.xls;*.xlsm")])
    ruta_excel_principal = ruta
    if not ruta:
        messagebox.showwarning("Aviso", "No se seleccionó archivo.")
        return None

    # Limpieza y tipos
    try:
        df = pd.read_excel(ruta, dtype={"Matrícula": str, "Grupo": str})
        df.dropna(subset=['Nombre completo'], inplace=True)
        df = df[df['Nombre completo'].astype(str).str.strip().str.lower() != 'nan']
        df = df[df['Nombre completo'].astype(str).str.strip() != '']
    except Exception as e:
        messagebox.showerror("Error Excel", f"No se pudo leer el archivo Excel.\n\n{e}")
        return None

    if df.empty:
        messagebox.showwarning("Aviso", "No se encontraron alumnos con nombre válido en el archivo.")
        return None

    datos = {col: df[col].astype(str).tolist() for col in df.columns}

    with open('datos.pkl', 'wb') as f:
        pickle.dump(datos, f)

    mostrar_alumnos(datos)
    return datos


def filtrar_datos(datos, matricula):
    datos_filtrados = {col: [] for col in datos.keys()}
    for i, m in enumerate(datos['Matrícula']):
        if m.strip() == matricula.strip():
            for col in datos:
                datos_filtrados[col].append(datos[col][i])
    return datos_filtrados


def cerrar_programa():
    sys.exit()


def mostrar_alumnos(datos):
    win = crear_ventana_toplevel("Seleccionar Alumno(s)", ancho=450, alto=600)
    win.protocol("WM_DELETE_WINDOW", cerrar_programa)

    ttk.Label(win, text="Seleccione uno o varios alumnos:", style="Title.TLabel").pack(pady=15)

    # Contenedor con scroll
    frame = ttk.Frame(win)
    frame.pack(fill="both", expand=True, padx=20, pady=5)

    canvas = tk.Canvas(frame, bg=COLOR_BG, highlightthickness=0)
    scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview, style="Vertical.TScrollbar")
    scrollable_frame = ttk.Frame(canvas)

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", width=380)
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Habilitar scroll con mouse
    def _on_mousewheel(event):
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    canvas.bind_all("<MouseWheel>", _on_mousewheel)

    # Lista de alumnos con checkboxes
    vars_checks = []
    for i, nombre in enumerate(datos['Nombre completo']):
        var = tk.BooleanVar(value=False)
        chk = ttk.Checkbutton(
            scrollable_frame,
            text=nombre,
            variable=var
        )
        chk.pack(anchor="w", padx=10, pady=2)
        vars_checks.append((var, i))

    # Procesar selección
    def procesar_seleccion():
        seleccionados = [i for var, i in vars_checks if var.get()]
        if not seleccionados:
            messagebox.showwarning("Atención", "Selecciona al menos un alumno.")
            return
        datos_seleccionados = {col: [] for col in datos.keys()}
        for col in datos.keys():
            for idx in seleccionados:
                datos_seleccionados[col].append(datos[col][idx])
        ventana_opciones(datos_seleccionados)

    # Botones de acción
    botones_frame = ttk.Frame(win)
    botones_frame.pack(pady=20)

    ttk.Button(
        botones_frame,
        text="✅ Procesar selección",
        style="Accent.TButton",
        command=procesar_seleccion
    ).pack(side="left", padx=10)

    def select_all():
        for var, _ in vars_checks:
            var.set(True)

    ttk.Button(
        botones_frame,
        text="☑ Seleccionar Todos",
        style="Secondary.TButton",
        command=select_all
    ).pack(side="left", padx=10)


def seleccionar_todos(datos, win=None):
    if win:
        pass
    ventana_opciones(datos)


def ventana_opciones(datos):
    """
    Ventana de opciones que procesa todos los alumnos seleccionados
    de forma secuencial al presionar los botones.
    """
    win = crear_ventana_toplevel("Opciones", ancho=450, alto=400)

    nombre_display = "todos los alumnos" if len(datos['Nombre completo']) > 1 else datos['Nombre completo'][0]
    
    ttk.Label(win, text=f"Opciones para:", style="Header.TLabel").pack(pady=(20, 5))
    ttk.Label(win, text=nombre_display, font=(FONT_FAMILY, 10, "italic")).pack(pady=(0, 20))

    # Crear expedientes para todos
    ttk.Button(
        win,
        text="📂 Crear Expedientes",
        style="Secondary.TButton",
        command=lambda: (crear_expedientes(datos), win.destroy())
    ).pack(pady=10, fill="x", padx=50)

    # Función auxiliar para procesar cada alumno uno por uno
    def procesar_actividad():
        for idx in range(len(datos["Nombre completo"])):
            alumno = {col: [datos[col][idx]] for col in datos.keys()}
            try:
                agregar_actividad_y_evidencia(alumno)
            except InterruptedError:
                break # Detener si el usuario cancela
        win.destroy()
    
    def procesar_consolidado_batch():
        # 1. Pedir Tutor UNICA VEZ
        tutor_global = seleccionar_tutor_gui("Selecciona Tutor para el CONSOLIDADO")
        if not tutor_global:
            return

        # 2. Loop
        for idx in range(len(datos["Nombre completo"])):
            alumno = {col: [datos[col][idx]] for col in datos.keys()}
            procesar_consolidado_alumno(alumno, tutor_global)
            
        messagebox.showinfo("Proceso Terminado", "Se ha procesado el Excel Consolidado para los alumnos seleccionados.")
        win.destroy()

    def procesar_entrevista():
        for idx in range(len(datos["Nombre completo"])):
            alumno = {col: [datos[col][idx]] for col in datos.keys()}
            agregar_entrevista_extendida(alumno)
        win.destroy()

    # Botones de actividad y entrevista
    ttk.Button(
        win,
        text="📝 Agregar actividad y evidencia (Semanal)",
        style="Accent.TButton",
        command=lambda: procesar_actividad()
    ).pack(pady=5, fill="x", padx=50)

    ttk.Button(
        win,
        text="📅 Procesar Consolidado (Todas las semanas)",
        style="Secondary.TButton",
        command=lambda: procesar_consolidado_batch()
    ).pack(pady=5, fill="x", padx=50)

    ttk.Button(
        win,
        text="💬 Agregar entrevista extendida",
        style="Warning.TButton", # Usamos estilo warning (naranja) si existiera, o default
        command=lambda: procesar_entrevista()
    ).pack(pady=5, fill="x", padx=50)

    # Botón cerrar
    ttk.Button(
        win,
        text="❌ Cerrar",
        style="Danger.TButton",
        command=win.destroy
    ).pack(pady=20, fill="x", padx=50)

# -------------------------------------------------------------------------------
# NUEVO: Implementación de "Agregar actividad y evidencia"
# -------------------------------------------------------------------------------

def procesar_consolidado_alumno(datos, tutor_global):
    """
    Busca al alumno en TODAS las hojas del Excel consolidado (ruta_excel_principal)
    e inserta los reportes encontrados en su documento Word.
    """
    global ruta_excel_principal
    global ruta_carpeta_global
    
    nombre = datos["Nombre completo"][0]
    nombre_may = nombre.upper()
    matricula = datos.get("Matrícula", [""])[0]
    programa = datos.get("Programa", [""])[0]
    
    print(f"Procesando CONSOLIDADO para: {nombre}")

    try:
        # 1. Obtener todas las hojas del Excel
        xl = pd.ExcelFile(ruta_excel_principal)
        hojas = xl.sheet_names # Lista de nombres de hojas
        
        # 2. Localizar Word
        ruta_alumno = localizar_carpeta_final_alumno(ruta_carpeta_global, matricula, nombre_may, programa)
        ruta_docx = encontrar_archivo_seguimiento(ruta_alumno)
        
        if not ruta_docx:
            print(f"  [Error] No hay Word para {nombre}")
            return

        # 3. Abrir Word UNA VEZ
        try:
            doc = docx.Document(ruta_docx)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir Word de {nombre}: {e}")
            return
            
        if not doc.tables:
            return

        tabla = doc.tables[0]
        cambios = False
        
        # 4. Iterar hojas
        for hoja in hojas:
            # Filtro básico: ignorar hojas que no parezcan fechas o reportes
            # Si el usuario dijo "fechas", asumimos que no son "Portada", "Datos", etc.
            # Intento leer la hoja usando la lógica existente
            try:
                # DEBUG_LEER_REPORTE=False para no spamear
                df_rep = leer_tabla_reporte(ruta_excel_principal, hoja=hoja)
            except ValueError:
                # Probablemente no tiene columnas Nombre/Apellido -> no es reporte
                continue
            except Exception:
                continue
                
            try:
                fila = filtrar_fila_reciente_por_alumno(df_rep, nombre)
                # Si encontramos fila, insertamos
                # Asumimos asistencia="Sí" si hay reporte, a menos que definamos lógica, 
                # pero para consolidado usaremos "Sí" por defecto (o N/A si fuera el caso).
                # Podríamos revisar el contenido para inferir "No", pero KISS por ahora.
                asistio = "Sí" 
                
                # Insertar (con chequeo de duplicados interno)
                inserto = _escribir_renglon_seguimiento(tabla, fila, tutor_global, asistio)
                if inserto:
                    cambios = True
                    print(f"  [Insertado] Hoja: {hoja}")
            except ValueError:
                # Alumno no está en esta hoja (semana)
                pass

        # 5. Guardar si hubo cambios
        if cambios:
            try:
                doc.save(ruta_docx)
                print(f"  [Guardado] {ruta_docx}")
            except PermissionError:
                messagebox.showerror("Error de Permiso", f"Cierra el archivo Word de {nombre}!")
        else:
            print("  [Sin cambios] No se agregaron nuevos registros (o duplicados).")
            
    except Exception as e:
        print(f"Error procesando {nombre}: {e}")
        # No bloquear todo el proceso por un alumno fallido
        pass


def agregar_actividad_y_evidencia(datos):
    """Flujo para un único alumno (Manual o Consolidado)."""
    global seleccion_primera_vez
    global ruta_carpeta_global
    global ruta_excel_principal

    try:
        # Preguntar si procesar SEMANA ACTUAL (Hoja 3) o CONSOLIDADO (Todas)
        # Esto podria ser una opción global, pero por ahora lo preguntamos al iniciar el proceso por lote o individual
        # Para no romper flujos existentes, si 'datos' viene de lote, deberíamos saber el modo.
        # Haremos una detección simple: Si el usuario selecciona "Procesar Consolidado" en el menú,
        # llamaremos a otra función. Si llama a esta, es el flujo MANUAL (Semana actual + imagen).
        pass
    except:
        pass
    
    # --- LOGICA ORIGINAL (Refactorizada) ---
    try:
        nombre = datos["Nombre completo"][0]
        nombre_may = nombre.upper()
        matricula = datos.get("Matrícula", [""])[0]
        programa = datos.get("Programa", [""])[0]

        if seleccion_primera_vez:
            ruta_carpeta_global = seleccionar_carpeta("Selecciona la carpeta MAESTRA:")
            ruta_excel_principal = seleccionar_excel_tabla_reporte()
            seleccion_primera_vez = False

        # 2) Leer reporte (Semana actual - Hoja index 2)
        df_rep = leer_tabla_reporte(ruta_excel_principal, hoja=2) # Default hoja 3
        fila = filtrar_fila_reciente_por_alumno(df_rep, nombre)

        # 3) Localizar doc
        ruta_alumno = localizar_carpeta_final_alumno(ruta_carpeta_global, matricula, nombre_may, programa)
        ruta_docx = encontrar_archivo_seguimiento(ruta_alumno)

        # 4) Insertar
        # Pide tutor/asistencia INTERACTIVAMENTE
        _, asistio = insertar_actividad_en_word(ruta_docx, fila, nombre_may)
        
        if asistio is None: return

        # 5) Evidencias
        carpeta_evid = encontrar_carpeta_evidencias(ruta_alumno)
        iniciales = obtener_iniciales(nombre)

        if asistio in ["Sí", "Tarea", "Falta justificada"]:
            
            # Avisos específicos
            if asistio == "Tarea":
                messagebox.showinfo("Aviso", "Por favor selecciona la imagen de la tarea.")
            elif asistio == "Falta justificada":
                messagebox.showinfo("Aviso", "Por favor selecciona la imagen del justificante.")

            titulo = f"Selecciona la evidencia para {nombre}"
            ruta_img = filedialog.askopenfilename(title=titulo, filetypes=[
                ("Imágenes", "*.jpg;*.jpeg;*.png;*.heic;*.webp;*.bmp")
            ])
            if ruta_img:
                actividad = simpledialog.askstring("Actividad", "Nombre de la actividad:")
                if actividad:
                    copiar_y_renombrar_evidencia(carpeta_evid, ruta_img, actividad, iniciales)
                    messagebox.showinfo("Listo", "Actividad agregada y evidencia guardada.")
            else:
                 messagebox.showinfo("Aviso", "No se seleccionó imagen.")
        
        elif asistio == "No":
            actividad = "BORRAR"
            num_sesion = proximo_numero_sesion(carpeta_evid, iniciales)
            actividad_limpia = "_".join(str(actividad).strip().split())
            nombre_dest = f"Sesión{num_sesion}_{actividad_limpia}_{iniciales}.jpeg"
            ruta_dest = os.path.join(carpeta_evid, nombre_dest)
            crear_imagen_blanca(ruta_dest)
            messagebox.showinfo("Listo", f"Creada evidencia en blanco: {nombre_dest}")

        elif asistio == "N/A":
            messagebox.showinfo("Listo", "Agregado N/A.")

    except Exception as e:
        messagebox.showerror("Error", f"Problema con {nombre}: {e}")


def agregar_entrevista_extendida(datos):
    """
    Flujo para agregar entrevista extendida: busca archivo de entrevista y anexa imágenes seleccionadas.
    """
    global seleccion_primera_vez
    global ruta_carpeta_global
    global ruta_excel_principal

    try:
        nombre = datos["Nombre completo"][0]
        nombre_may = nombre.upper()
        matricula = datos.get("Matrícula", [""])[0]
        programa = datos.get("Programa", [""])[0]

        if seleccion_primera_vez:

            # 1) Seleccionar carpeta maestra y Excel semanal
            ruta_carpeta_global = seleccionar_carpeta("Selecciona la carpeta MAESTRA que contiene todas las carpetas de alumnos:")
            #ruta_carpeta_global = ruta_carpeta_global
            ruta_excel_principal = seleccionar_excel_tabla_reporte() #Validar que si se haya seleccionado un excel
            seleccion_primera_vez = False

        # 1) Ubicar carpeta final del alumno y archivo de entrevista
        ruta_alumno = localizar_carpeta_final_alumno(ruta_carpeta_global, matricula, nombre_may, programa)
        ruta_docx = encontrar_archivo_entrevista(ruta_alumno)

        if not ruta_docx:
            messagebox.showerror("Error", f"No se encontró archivo de entrevista para {nombre}")
            return

        # 2) Seleccionar imágenes
        rutas_imagenes = filedialog.askopenfilenames(
            title=f"Selecciona las imágenes de la entrevista de {nombre}",
            filetypes=[("Archivos de imagen", "*.jpg *.jpeg *.png *.bmp *.tiff")]
        )
        if not rutas_imagenes:
            return

        # Ordenar imágenes por fecha de modificación (más antigua primero)
        rutas_imagenes = sorted(rutas_imagenes, key=os.path.getmtime)

        # 3) Insertar imágenes en el documento
        try:
            doc = Document(ruta_docx)
        except Exception as e:
            messagebox.showerror("Error Archivo", f"No se pudo abrir el archivo Word.\n{e}")
            return

        section = doc.sections[0]

        # Dimensiones de página
        ancho_pagina = section.page_width.inches
        alto_pagina = section.page_height.inches
        margen_izq = section.left_margin.inches
        margen_der = section.right_margin.inches
        margen_sup = section.top_margin.inches
        margen_inf = section.bottom_margin.inches

        ancho_max = ancho_pagina - (margen_izq + margen_der)
        alto_max = alto_pagina - (margen_sup + margen_inf)

        for i, img_path in enumerate(rutas_imagenes):
            if i > 0:
                # Insertar salto de página limpio antes de la imagen
                doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
                espacio_disp = alto_max
            else:
                # Primera imagen: reservar espacio por texto inicial
                espacio_disp = alto_max - 2.0
                if espacio_disp < alto_max * 0.5:
                    espacio_disp = alto_max * 0.5

            # Abrir imagen y calcular escalado
            with Image.open(img_path) as im:
                w, h = im.size
                dpi = im.info.get("dpi", (96, 96))
                w_in = w / dpi[0]
                h_in = h / dpi[1]

                factor = min(ancho_max / w_in, espacio_disp / h_in)
                ancho_final = Inches(w_in * factor)
                alto_final = Inches(h_in * factor)

            # Insertar imagen
            doc.add_picture(img_path, width=ancho_final, height=alto_final)

        # 4) Guardar cambios
        try:
            doc.save(ruta_docx)
            messagebox.showinfo("Éxito", f"Entrevista actualizada con evidencias para {nombre}")
        except PermissionError:
            messagebox.showerror("Error de Permiso", "No se pudo guardar el archivo. ¡Ciérralo si está abierto!")

    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un problema al generar las entrevistas: {e}")

def previsualizar_datos_excel(datos_fila):
    """
    Muestra una ventana con los datos que se van a insertar en Word.
    `datos_fila` debe ser un diccionario con las claves: 'Fecha', 'Tema', 'Situación'.
    """
    ventana = crear_ventana_toplevel("Previsualización de datos (DEBUG)", ancho=400, alto=250)
    
    ttk.Label(ventana, text="Datos a insertar en Word:", style="Header.TLabel").pack(pady=10)

    ttk.Label(ventana, text=f"Fecha de atención: {datos_fila['Fecha']}").pack(pady=5)
    ttk.Label(ventana, text=f"Tema o asunto tratado: {datos_fila['Tema']}").pack(pady=5)
    ttk.Label(ventana, text=f"Situación del alumno: {datos_fila['Situación']}").pack(pady=5)
    
    ttk.Button(ventana, text="OK, continuar", command=ventana.destroy).pack(pady=15)
    ventana.grab_set()
    ventana.mainloop()

# -------------------------------------------------------------------------------
# Utilidades ya existentes para reemplazos y fechas
# -------------------------------------------------------------------------------

dias_semana = {"lunes": 0, "martes": 1, "miércoles": 2, "jueves": 3, "viernes": 4, "sábado": 5, "domingo": 6}


def proxima_reunion(dia, hora_inicio):
    if not dia or pd.isna(dia) or str(dia).lower() == 'nan':
        return "FECHA NO DEFINIDA"
    hoy = datetime.now()
    dia_actual = hoy.weekday()
    hora_act = hoy.time()
    partes = str(hora_inicio).split(":")
    try:
        h = int(partes[0])
        m = int(partes[1])
    except (ValueError, IndexError):
        h, m = 0, 0
    hora_ini = time(h, m)
    dia_reu = dias_semana.get(str(dia).lower().strip(), 0)
    if dia_reu < dia_actual or (dia_reu == dia_actual and hora_ini <= hora_act):
        dias_para = 7 - (dia_actual - dia_reu)
    else:
        dias_para = dia_reu - dia_actual
    prox = hoy + timedelta(days=dias_para)
    return f"{str(dia).upper()} {prox.day} DE {prox.strftime('%B').upper()} A LAS {hora_inicio}"


def formatear_hora(valor):
    if isinstance(valor, (datetime, time)):
        return valor.strftime("%H:%M")
    try:
        valor_str = str(valor)
        if ":" in valor_str:
            partes = valor_str.split(":")
            if len(partes) >= 2:
                return f"{partes[0].zfill(2)}:{partes[1].zfill(2)}"
        float_val = float(valor)
        frac_dia = float_val * 24
        horas = int(frac_dia)
        minutos = int((frac_dia - horas) * 60)
        return f"{horas:02d}:{minutos:02d}"
    except (ValueError, TypeError):
        return "00:00"


def reemplazar_en_docx(doc, markers):
    for para in doc.paragraphs:
        for mk, txt in markers.items():
            if mk in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if mk in inline[i].text:
                        text = inline[i].text.replace(mk, str(txt))
                        inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for mk, txt in markers.items():
                        if mk in para.text:
                            inline = para.runs
                            for i in range(len(inline)):
                                if mk in inline[i].text:
                                    text = inline[i].text.replace(mk, str(txt))
                                    inline[i].text = text


# -------------------------------------------------------------------------------
# App Tkinter
# -------------------------------------------------------------------------------

def ejecutar_y_ocultar(func, win):
    win.withdraw()
    func()


def iniciar_app():
    global root
    root = tk.Tk()
    root.title("Gestión de Expedientes")
    root.geometry("300x200")
    
    # Aplicar tema oscuro
    aplicar_estilo(root)
    
    root.withdraw()
    ventana_bienvenida()
    root.mainloop()


def ventana_bienvenida():
    win = crear_ventana_toplevel("Bienvenida", ancho=450, alto=300)

    # Mensaje de bienvenida
    ttk.Label(
        win, 
        text="Bienvenido al Gestor de Expedientes", 
        style="Title.TLabel"
    ).pack(pady=(40, 20))

    # Botón destacado
    ttk.Button(
        win, 
        text="🚀 Iniciar Aplicación", 
        style="Accent.TButton",
        command=lambda: ejecutar_y_ocultar(cargar_datos_excel, win)
    ).pack(pady=20, ipadx=10, ipady=5)


if __name__ == "__main__":
    iniciar_app()
